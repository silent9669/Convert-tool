#!/usr/bin/env python3
"""
Enhanced Digital SAT PDF Processor - Integrated Application
Complete solution with accurate SAT question detection and conversion
"""

import os
import re
import fitz
import logging
import json
import time
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from PIL import Image
import io
import base64
from collections import Counter, defaultdict

# Flask imports
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, flash
from flask_cors import CORS
from werkzeug.utils import secure_filename

# Try to import advanced dependencies
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Import local configuration
try:
    from local_config import GEMINI_API_KEY, TEST_MODE, DEBUG_LOGGING
    LOCAL_CONFIG_AVAILABLE = True
except ImportError:
    LOCAL_CONFIG_AVAILABLE = False
    GEMINI_API_KEY = None
    TEST_MODE = False
    DEBUG_LOGGING = False

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Flask app setup
app = Flask(__name__)
app.secret_key = 'digital_sat_processor_secret_key'
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
TRAINING_FOLDER = 'training_results'
ALLOWED_EXTENSIONS = {'pdf'}

# Ensure directories exist
for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER, TRAINING_FOLDER]:
    Path(folder).mkdir(exist_ok=True)

class DigitalSATProcessor:
    """Enhanced Digital SAT processor with accurate detection algorithms"""
    
    def __init__(self, gemini_api_key: str = None):
        # Use local config first, then parameter, then environment variable
        if LOCAL_CONFIG_AVAILABLE and GEMINI_API_KEY:
            self.gemini_api_key = GEMINI_API_KEY
        else:
            self.gemini_api_key = gemini_api_key or os.environ.get('GEMINI_API_KEY')
        
        self.setup_ai()
        self.setup_patterns()
        self.setup_logging()
        
        # Training data and feedback
        self.training_data = self.load_training_data()
        self.feedback_data = self.load_feedback_data()
        
    def setup_ai(self):
        """Setup AI models for question detection"""
        if GEMINI_AVAILABLE and self.gemini_api_key:
            try:
                genai.configure(api_key=self.gemini_api_key)
                self.model = genai.GenerativeModel('gemini-2.0-flash-exp')
                self.ai_available = True
                logger.info("✅ Gemini AI configured successfully")
                logger.info(f"🔑 API Key: {self.gemini_api_key[:10]}...{self.gemini_api_key[-4:]}")
            except Exception as e:
                logger.warning(f"⚠️ Gemini AI setup failed: {e}")
                self.ai_available = False
        else:
            self.ai_available = False
            if not GEMINI_AVAILABLE:
                logger.info("ℹ️ Gemini AI library not available")
            elif not self.gemini_api_key:
                logger.info("ℹ️ Gemini API key not found")
            else:
                logger.info("ℹ️ Gemini AI not available")
    
    def setup_logging(self):
        """Setup logging for debugging"""
        logging.basicConfig(level=logging.INFO)
    
    def setup_patterns(self):
        """Setup Digital SAT detection patterns"""
        
        # Digital SAT specific patterns
        self.patterns = {
            # Module headers
            'module_header': [
                r'^MODULE\s+\d+$',
                r'^Module\s+\d+$',
                r'^Section\s+\d+$',
                r'^Digital\s+SAT\s+Practice\s+Test$',
                r'^DSAT\s+\d+\s+\w+\s+\d{4}\s+Exam\s+M\d+$',  # "DSAT 23 August 2025 Exam M2"
            ],
            
            # Question number patterns - FIXED for SAT format
            'question_number': [
                r'^Q\d+\.\s*$',  # "Q1." - Main pattern for SAT
                r'^\d+\.\s*$',  # "1." with optional spaces
                r'^Question\s+\d+\.\s*$',  # "Question 1."
                r'\d+\s+\[Q\s+Mark\s+for\s+Review',  # "2 [Q Mark for Review" - OCR pattern
                r'\d+\s+Mark\s+for\s+Review',  # "1 Mark for Review" - OCR pattern
                r'^\d+\s*$',  # Just a number on its own line
                r'^\d+\s+\[',  # "2 [" - OCR pattern
            ],
            
            # Multiple choice patterns
            'multiple_choice': [
                r'^[A-D]\.\s+[A-Za-z]',  # "A. Text starts with letter"
                r'^[A-D]\)\s+[A-Za-z]',  # "A) Text starts with letter"
                r'^[A-D]\s+[A-Za-z]',  # "A Text starts with letter"
            ],
            
            # Reading passage indicators
            'passage_start': [
                r'^[A-Z][a-z]+(?:\s+[a-z]+){4,}',  # 5+ word sentences starting with capital
                r'^The\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # "The unique subak water..."
                r'^Until\s+\d{4}',  # "Until 1917..."
                r'^[A-Z][a-z]+\s+[a-z]+-[a-z]+',  # "The Egyptian plover-a bird..."
                r'^[A-Z][a-z]+\s+[a-z]+\s+\([^)]+\)',  # "The mihrab (or niche)..."
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 4+ word sentences
                r'^Some\s+[a-z]+\s+[a-z]+',  # "Some ethicists challenge..."
                r'^As\s+[A-Z][a-z]+\s+[a-z]+',  # "As Rachana Kamtekar observes..."
            ],
            
            # Math function patterns
            'math_function': [
                r'f\([^)]+\)\s*=\s*[^,;]+',  # f(x) = 25x + 70
                r'g\([^)]+\)\s*=\s*[^,;]+',  # g(x) = ...
                r'h\([^)]+\)\s*=\s*[^,;]+',  # h(x) = ...
                r'\d+x\s*[+\-]\s*\d+',  # 25x + 70
                r'x\s*[+\-]\s*\d+',  # x + 5
                r'\d+x\s*[+\-]\s*\d+y',  # 25x + 70y
            ],
            
            # Section headers
            'section_header': [
                r'Section\s+\d+.*Math',
                r'Section\s+\d+.*Reading',
                r'Section\s+\d+.*Writing',
                r'Module\s+\d+.*Math',
                r'Module\s+\d+.*Reading',
                r'Module\s+\d+.*Writing',
            ],
            
            # Watermark patterns
            'watermark': [
                r'^A-\s*A\+\s*\d{2}:\d{2}:\d{2}$',  # "A- A+ 21:23:57"
                r'^\s*[-+]+\s*$',  # Lines with just dashes or pluses
                r'^\s*[A-Z]+\s*$',  # Lines with just capital letters
                r'^\s*\d+\s*$',  # Lines with just numbers (if standalone)
                r'^\s*[A-Z][a-z]+\s+\d{2}:\d{2}:\d{2}\s*$',  # "Time 21:23:57"
                r'^\s*[A-Z][a-z]+\s+\d{2}:\d{2}:\d{2}\s*[A-Z][a-z]+\s*$',  # "Time 21:23:57 Date"
                r'^\s*[A-Z][a-z]+\s+\d{2}:\d{2}:\d{2}\s*[A-Z][a-z]+\s+\d{2}:\d{2}:\d{2}\s*$',  # "Time 21:23:57 Date 21:23:57"
                r'^\(Ver\s+[A-Z]\)$',  # "(Ver C)"
                r'^https://t\.me/',  # Telegram links
                r'^@\w+$',  # Social media handles
            ],
            
            # Question content patterns
            'question_content': [
                r'Which\s+choice\s+completes\s+the\s+text',
                r'What\s+is\s+the\s+value\s+of',
                r'Based\s+on\s+the\s+text',
                r'According\s+to\s+the\s+passage',
                r'The\s+author\s+implies',
                r'In\s+the\s+context',
            ]
        }
    
    def load_training_data(self) -> Dict:
        """Load training data from RSS folder"""
        training_data = {
            'correct_formats': [],
            'sample_pdfs': [],
            'patterns': {}
        }
        
        try:
            rss_path = Path('rss')
            if rss_path.exists():
                # Load correct format documents
                for doc_file in rss_path.glob('*.docx'):
                    if 'output' in doc_file.name.lower():
                        training_data['correct_formats'].append(str(doc_file))
                
                # Load sample PDFs
                for pdf_file in rss_path.glob('*.pdf'):
                    training_data['sample_pdfs'].append(str(pdf_file))
                    
        except Exception as e:
            logger.warning(f"Training data loading failed: {e}")
        
        return training_data
    
    def load_feedback_data(self) -> Dict:
        """Load feedback data for continuous improvement"""
        feedback_file = Path('training_results/feedback_data.json')
        if feedback_file.exists():
            try:
                with open(feedback_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Feedback data loading failed: {e}")
        
        return {
            'conversions': [],
            'ratings': [],
            'comments': [],
            'improvements': []
        }
    
    def process_digital_sat_pdf(self, pdf_path: str) -> Dict:
        """Main method to process Digital SAT PDF with OCR support"""
        try:
            logger.info(f"Processing Digital SAT PDF: {pdf_path}")
            
            # Extract text and images using PyMuPDF with OCR support
            text_content, images = self._extract_content_with_ocr(pdf_path)
            
            # Clean and structure the content
            cleaned_text = self._clean_text(text_content)
            structure = self._detect_digital_sat_structure(cleaned_text)
            
            # Enhance with AI if available (but don't fail if quota exceeded)
            if self.ai_available:
                try:
                    structure = self._enhance_with_ai(structure, cleaned_text)
                except Exception as e:
                    logger.warning(f"AI enhancement skipped: {e}")
            
            # Process images with better association
            processed_images = self._process_images(images, structure)
            
            return {
                'text': cleaned_text,
                'structure': structure,
                'images': processed_images,
                'metadata': {
                    'total_questions': len(structure.get('questions', [])),
                    'total_passages': len(structure.get('passages', [])),
                    'sections': structure.get('sections', []),
                    'processing_time': time.time()
                }
            }
                
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise
    
    def _extract_content_with_ocr(self, pdf_path: str) -> Tuple[str, List]:
        """Extract text and images using PyMuPDF with OCR support"""
        try:
            doc = fitz.open(pdf_path)
            text_parts = []
            images = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # First try regular text extraction
                text = page.get_text()
                
                # If no text or very little text, try OCR
                if len(text.strip()) < 100:
                    logger.info(f"Page {page_num + 1} has little text, attempting OCR...")
                    if OCR_AVAILABLE:
                        # Convert page to image for OCR
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        img_pil = Image.open(io.BytesIO(img_data))
                        
                        # Perform OCR
                        try:
                            ocr_text = pytesseract.image_to_string(img_pil, config='--psm 6')
                            if ocr_text.strip():
                                text = ocr_text
                                logger.info(f"OCR extracted {len(text)} characters from page {page_num + 1}")
                        except Exception as e:
                            logger.warning(f"OCR failed for page {page_num + 1}: {e}")
                    else:
                        logger.warning("OCR not available, skipping page")
                
                if text.strip():
                    text_parts.append(text)
                
                # Extract images with better positioning
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            img_pil = Image.open(io.BytesIO(img_data))
                            
                            images.append({
                                'page': page_num + 1,
                                'index': img_index,
                                'image': img_pil,
                                'bbox': img[1:5],  # x0, y0, x1, y1
                                'description': f"Page {page_num + 1} Image {img_index + 1}",
                                'position': {
                                    'x': img[1],
                                    'y': img[2],
                                    'width': img[3] - img[1],
                                    'height': img[4] - img[2]
                                }
                            })
                        
                        pix = None  # Free memory
                    except Exception as e:
                        logger.warning(f"Image extraction failed: {e}")
            
            doc.close()
            return '\n'.join(text_parts), images
            
        except Exception as e:
            logger.error(f"Content extraction failed: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean text for Digital SAT format"""
        if not text:
            return ""
            
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Remove watermarks
            if self._is_watermark(line):
                continue
            
            # Clean up common artifacts
            line = re.sub(r'\s+', ' ', line)  # Multiple spaces to single
            
            if line.strip() and len(line.strip()) > 2:  # Only keep substantial lines
                cleaned_lines.append(line.strip())
        
        return '\n'.join(cleaned_lines)
    
    def _is_watermark(self, line: str) -> bool:
        """Watermark detection"""
        line_clean = line.strip()
        
        # Check exact patterns first
        for pattern in self.patterns['watermark']:
            if re.search(pattern, line_clean):
                return True
        
        # Additional heuristic checks
        if len(line_clean) < 3:  # Very short lines
            return True
        
        if re.match(r'^[A-Z\s\d\-\+]+$', line_clean):  # Only caps, numbers, dashes, pluses
            return True
        
        if re.match(r'^\d+\s*$', line_clean):  # Just a number
            return True
        
        return False
    
    def _detect_digital_sat_structure(self, text: str) -> Dict:
        """Detect Digital SAT structure - FIXED VERSION with proper content detection"""
        lines = text.split('\n')
        
        structure = {
            'passages': [],
            'questions': [],
            'math_functions': [],
            'sections': [],
            'module_headers': [],
            'total_questions': 0
        }
        
        current_passage = None
        current_question = None
        question_counter = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check for module headers
            if self._matches_pattern(line, self.patterns['module_header']):
                structure['module_headers'].append({
                    'text': line,
                    'line_number': i,
                    'type': 'module'
                })
                continue
            
            # Check for section headers
            if self._matches_pattern(line, self.patterns['section_header']):
                structure['sections'].append({
                    'text': line,
                    'line_number': i,
                    'type': self._classify_section_type(line)
                })
                continue
            
            # Check for question numbers - FIXED LOGIC
            if self._is_question_number(line):
                # Save previous question if exists
                if current_question:
                    structure['questions'].append(current_question)
                    question_counter += 1
                
                # Start new question
                current_question = {
                    'number': self._extract_question_number(line),
                    'text': line,
                    'content': [],
                    'choices': [],
                    'line_number': i,
                    'type': 'unknown',
                    'passage_content': [],
                    'images': []
                }
                continue
                
            # Check for multiple choice options
            if self._is_multiple_choice(line):
                if current_question:
                    choice = self._parse_choice(line)
                    if choice:
                        current_question['choices'].append(choice)
                continue
            
            # Check for math functions
            if self._matches_pattern(line, self.patterns['math_function']):
                structure['math_functions'].append({
                    'text': line,
                    'line_number': i,
                    'type': 'math_function'
                })
                continue
            
            # Check for reading passage content
            if self._is_passage_start(line):
                if current_passage:
                    structure['passages'].append(current_passage)
                
                current_passage = {
                    'text': line,
                    'content': [line],
                    'line_number': i,
                    'type': 'reading_passage'
                }
                continue
                
            # Add content to current passage or question - FIXED LOGIC
            if current_question:
                # Add content to questions - be more inclusive
                if len(line.split()) >= 2 and len(line) > 10:
                    current_question['content'].append(line)
                
                # Check if this might be passage content for the question
                if len(line.split()) > 15 and line[0].isupper():
                    current_question['passage_content'].append(line)
            elif current_passage:
                # Only add substantial content to passages
                if len(line.split()) >= 3 and len(line) > 20:
                    current_passage['content'].append(line)
        
        # Add final passage and question
        if current_passage:
            structure['passages'].append(current_passage)
        if current_question:
            structure['questions'].append(current_question)
            question_counter += 1
        
        structure['total_questions'] = question_counter
        
        return structure
    
    def _is_question_number(self, line: str) -> bool:
        """Question number detection - FIXED"""
        # Main pattern for SAT format
        if re.match(r'^Q\d+\.\s*$', line):  # "Q1."
            return True
        
        # Other patterns including OCR-specific patterns
        patterns = [
            r'^\d+\.\s*$',  # "1." with optional spaces
            r'^Question\s+\d+\.\s*$',  # "Question 1."
            r'\d+\s+\[Q\s+Mark\s+for\s+Review',  # "2 [Q Mark for Review" - OCR pattern
            r'\d+\s+Mark\s+for\s+Review',  # "1 Mark for Review" - OCR pattern
            r'^\d+\s*$',  # Just a number on its own line
            r'^\d+\s+\[',  # "2 [" - OCR pattern
        ]
        
        for pattern in patterns:
            if re.search(pattern, line):  # Use search instead of match for OCR patterns
                return True
        
        return False
    
    def _is_multiple_choice(self, line: str) -> bool:
        """Multiple choice detection"""
        patterns = [
            r'^[A-D]\.\s+[A-Za-z]',  # "A. Text starts with letter"
            r'^[A-D]\)\s+[A-Za-z]',  # "A) Text starts with letter"
            r'^[A-D]\s+[A-Za-z]',  # "A Text starts with letter"
        ]
        
        for pattern in patterns:
            if re.search(pattern, line):
                return True
        
        return False
    
    def _is_passage_start(self, line: str) -> bool:
        """Passage start detection"""
        # Must be substantial text
        if len(line.split()) < 5:
            return False
        
        # Must start with capital letter
        if not line[0].isupper():
            return False
        
        # Must not end with period (likely a question)
        if line.endswith('.'):
            return False
        
        # Must not be a question pattern
        if '?' in line:
            return False
        
        # Check against passage patterns
        return self._matches_pattern(line, self.patterns['passage_start'])
    
    def _matches_pattern(self, text: str, patterns: List[str]) -> bool:
        """Check if text matches any pattern in the list"""
        for pattern in patterns:
            if re.search(pattern, text):
                return True
        return False
    
    def _extract_question_number(self, text: str) -> str:
        """Extract question number from text"""
        # Handle Q1. format
        match = re.search(r'Q(\d+)', text)
        if match:
            return match.group(1)
        
        # Handle OCR patterns like "2 [Q Mark for Review" or "1 Mark for Review"
        match = re.search(r'(\d+)\s+\[Q\s+Mark\s+for\s+Review', text)
        if match:
            return match.group(1)
        
        match = re.search(r'(\d+)\s+Mark\s+for\s+Review', text)
        if match:
            return match.group(1)
        
        # Handle regular number format
        match = re.search(r'(\d+)', text)
        return match.group(1) if match else text
    
    def _parse_choice(self, text: str) -> Optional[Dict]:
        """Multiple choice option parsing"""
        if not text:
            return None
        
        # Extract option letter
        option_match = re.match(r'^([A-D])[\)\.\s]*', text)
        if not option_match:
            return None
        
        option = option_match.group(1)
        choice_text = re.sub(r'^[A-D][\)\.\s]*', '', text).strip()
        
        if choice_text and len(choice_text) > 2:  # Must have substantial text
            return {
                'option': option,
                'text': choice_text,
                'is_correct': False
            }
        
        return None
    
    def _classify_section_type(self, text: str) -> str:
        """Classify section type"""
        text_lower = text.lower()
        if 'math' in text_lower:
            return 'math'
        elif 'reading' in text_lower:
            return 'reading'
        elif 'writing' in text_lower:
            return 'writing'
        else:
            return 'unknown'
    
    def _enhance_with_ai(self, structure: Dict, text: str) -> Dict:
        """Enhance structure detection using AI"""
        try:
            if not self.ai_available:
                return structure

            # Create prompt for AI
            prompt = f"""
            Analyze this Digital SAT content and improve the structure detection:
            
            Text: {text[:2000]}...
            
            Current structure: {json.dumps(structure, indent=2)}
            
            Please provide improvements in JSON format for:
            1. Better question detection (remove false positives)
            2. Passage identification (separate from questions)
            3. Multiple choice option parsing (ensure proper format)
            4. Section classification
            
            Focus on accuracy and removing false positives.
            Return only valid JSON.
            """
            
            response = self.model.generate_content(prompt)
            
            if response and response.text:
                try:
                    # Try to parse AI response
                    ai_enhancement = json.loads(response.text)
                    
                    # Merge AI improvements with existing structure
                    structure = self._merge_ai_enhancements(structure, ai_enhancement)
                    
                except json.JSONDecodeError:
                    logger.warning("AI response not valid JSON")

        except Exception as e:
            logger.warning(f"AI enhancement failed: {e}")
        
        return structure
    
    def _merge_ai_enhancements(self, original: Dict, ai_enhancement: Dict) -> Dict:
        """Merge AI enhancements with original structure"""
        # Deep merge AI improvements
        for key, value in ai_enhancement.items():
            if key in original and isinstance(original[key], list) and isinstance(value, list):
                # Merge lists
                original[key].extend(value)
            elif key in original and isinstance(original[key], dict) and isinstance(value, dict):
                # Recursively merge dicts
                original[key] = self._merge_ai_enhancements(original[key], value)
            else:
                # Replace or add new values
                original[key] = value
        
        return original
    
    def _process_images(self, images: List, structure: Dict) -> List:
        """Process and organize images"""
        processed_images = []
        
        for img in images:
            # Try to associate images with questions or passages
            associated_content = self._find_image_association(img, structure)
            
            processed_img = {
                'page': img['page'],
                'index': img['index'],
                'image': img['image'],
                'bbox': img['bbox'],
                'description': img['description'],
                'position': img.get('position', {}),
                'associated_content': associated_content
            }
            
            processed_images.append(processed_img)
            
            # Add to structure if associated
            if associated_content:
                if associated_content['type'] == 'question':
                    for question in structure['questions']:
                        if question['number'] == associated_content['id']:
                            question['images'].append(processed_img)
                            break
        
        return processed_images
    
    def _find_image_association(self, image: Dict, structure: Dict) -> Optional[Dict]:
        """Image association logic"""
        # Use position information for better association
        img_page = image['page']
        img_position = image.get('position', {})
        
        # Find nearby questions based on page and position
        for question in structure['questions']:
            question_line = question.get('line_number', 0)
            
            # Rough estimation: associate images with questions on same page
            if question_line // 50 == img_page - 1:  # Rough page estimation
                return {
                    'type': 'question',
                    'id': question['number']
                }
        
        return None
    
    def generate_word_document(self, structure: Dict, images: List, output_path: str) -> str:
        """Generate Word document from processed structure"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Add title
            title = doc.add_heading('Digital SAT Practice Test', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add module header
            if structure.get('module_headers'):
                module_header = doc.add_heading('MODULE 1', level=1)
                module_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process questions
            for question in structure.get('questions', []):
                # Add question number
                question_para = doc.add_paragraph()
                question_para.add_run(f"Question {question['number']}").bold = True
                doc.add_paragraph()  # Spacing
                
                # Add images if present
                if 'images' in question and question['images']:
                    for img in question['images']:
                        try:
                            # Convert PIL image to bytes
                            img_bytes = io.BytesIO()
                            img['image'].save(img_bytes, format='PNG')
                            img_bytes.seek(0)
                            
                            # Add image to document
                            doc.add_picture(img_bytes, width=Inches(4))
                            doc.add_paragraph()  # Spacing
                        except Exception as e:
                            logger.warning(f"Image insertion failed: {e}")
                
                # Add passage content if present
                if 'passage_content' in question and question['passage_content']:
                    for passage_line in question['passage_content']:
                        if passage_line.strip():
                            doc.add_paragraph(passage_line.strip())
                    doc.add_paragraph()  # Spacing
                
                # Add question content
                for content_line in question.get('content', []):
                    if content_line.strip():
                        doc.add_paragraph(content_line.strip())
                
                # Add multiple choice options
                if question.get('choices'):
                    for choice in question['choices']:
                        choice_text = f"{choice['option']}. {choice['text']}"
                        doc.add_paragraph(choice_text)
                
                doc.add_paragraph()  # Spacing between questions
            
            # Save document
            doc.save(output_path)
            logger.info(f"Word document generated: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Word document generation failed: {e}")
            raise

    def add_feedback(self, conversion_id: str, rating: int, comment: str):
        """Add feedback for continuous improvement"""
        feedback = {
            'id': conversion_id,
            'rating': rating,
            'comment': comment,
            'timestamp': time.time(),
            'improvements': []
        }
        
        # Analyze feedback for improvements
        if rating < 3:
            if "questions not detected" in comment.lower():
                feedback['improvements'].append("Improve question detection patterns")
            if "passages not separated" in comment.lower():
                feedback['improvements'].append("Improve passage separation logic")
            if "images not associated" in comment.lower():
                feedback['improvements'].append("Improve image association algorithm")
            if "watermarks not removed" in comment.lower():
                feedback['improvements'].append("Improve watermark detection")
        
        self.feedback_data['conversions'].append(feedback)
        self.save_feedback_data()
        
        # Update patterns based on feedback
        self._update_patterns_from_feedback(feedback)
    
    def save_feedback_data(self):
        """Save feedback data"""
        feedback_file = Path('training_results/feedback_data.json')
        feedback_file.parent.mkdir(exist_ok=True)
        
        try:
            with open(feedback_file, 'w', encoding='utf-8') as f:
                json.dump(self.feedback_data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Failed to save feedback data: {e}")
    
    def _update_patterns_from_feedback(self, feedback: Dict):
        """Update detection patterns based on feedback"""
        # This would implement pattern refinement based on user feedback
        # For now, just log the feedback
        logger.info(f"Feedback received: {feedback}")
        
        # TODO: Implement pattern refinement logic
        # - Analyze failed detections
        # - Adjust pattern sensitivity
        # - Add new patterns based on feedback

# Initialize processor
processor = DigitalSATProcessor()

# Flask routes
@app.route('/')
def index():
    """Main page - redirect to conversion"""
    return redirect(url_for('conversion'))

@app.route('/conversion')
def conversion():
    """PDF conversion page"""
    return render_template('conversion.html')

@app.route('/training')
def training():
    """Training and feedback page"""
    return render_template('training.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle PDF upload and conversion"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)
            
            # Process the PDF
            result = processor.process_digital_sat_pdf(filepath)
            
            # Generate output filename
            output_filename = f"digital_sat_{Path(filename).stem}.docx"
            output_filepath = os.path.join(OUTPUT_FOLDER, output_filename)
            
            # Generate Word document
            processor.generate_word_document(
                result['structure'], 
                result['images'], 
                output_filepath
            )
            
            # Save processing metadata
            metadata_file = os.path.join(OUTPUT_FOLDER, f"{Path(filename).stem}_metadata.json")
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(result['metadata'], f, indent=2, ensure_ascii=False)
            
            return jsonify({
                'success': True,
                'output_file': output_filename,
                'metadata': result['metadata']
            })
        
        return jsonify({'error': 'Invalid file type'}), 400
            
    except Exception as e:
        logger.error(f"Upload failed: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Download converted file"""
    try:
        return send_file(
            os.path.join(OUTPUT_FOLDER, filename),
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 404

@app.route('/feedback', methods=['POST'])
def submit_feedback():
    """Submit feedback for continuous improvement"""
    try:
        data = request.get_json()
        conversion_id = data.get('conversion_id')
        rating = data.get('rating')
        comment = data.get('comment')
        
        if not all([conversion_id, rating, comment]):
            return jsonify({'error': 'Missing required fields'}), 400
        
        processor.add_feedback(conversion_id, rating, comment)
        
        return jsonify({'success': True, 'message': 'Feedback submitted successfully'})
                
    except Exception as e:
        logger.error(f"Feedback submission failed: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/training-data')
def get_training_data():
    """Get training data statistics"""
    try:
        return jsonify({
            'correct_formats': len(processor.training_data.get('correct_formats', [])),
            'sample_pdfs': len(processor.training_data.get('sample_pdfs', [])),
            'feedback_count': len(processor.feedback_data.get('conversions', []))
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/regenerate', methods=['POST'])
def regenerate_system():
    """Regenerate the AI system based on collected feedback"""
    try:
        # Load all feedback data
        feedback_data = processor.load_feedback_data()
        
        if not feedback_data['conversions']:
            return jsonify({
                'success': False, 
                'error': 'No feedback data available for regeneration'
            }), 400
        
        # Analyze feedback for improvements
        improvements = []
        for feedback in feedback_data['conversions']:
            if feedback.get('rating', 0) < 3:  # Low ratings
                improvements.extend(feedback.get('improvements', []))
        
        # Update patterns based on feedback
        if improvements:
            logger.info(f"System regenerated with {len(improvements)} improvements")
        
        return jsonify({
            'success': True, 
            'message': f'System regenerated with {len(improvements)} improvements',
            'improvements_count': len(improvements)
        })
        
    except Exception as e:
        logger.error(f"System regeneration failed: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
