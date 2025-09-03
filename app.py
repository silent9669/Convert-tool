#!/usr/bin/env python3
"""
Enhanced PDF Watermark Remover with Word Conversion
Two sections: English (text processing) and Math (LaTeX conversion)
Based on advanced watermark removal techniques from multiple repositories
Optimized for Python 3.11.9 with enhanced detection capabilities
"""

import os
import sys
import logging
import traceback
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import time
import re
from PIL import Image
import io
from advanced_sat_processor import AdvancedSATProcessor

# Check Python version
if sys.version_info < (3, 11):
    print("Error: Python 3.11+ is required. Current version:", sys.version)
    sys.exit(1)

print(f"Python {sys.version} detected - proceeding with installation...")

try:
    from flask import Flask, request, jsonify, send_file, render_template_string, redirect
    from flask_cors import CORS
    from werkzeug.utils import secure_filename
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    import io
    from PIL import Image
    
    # Try to import OCR support (optional)
    try:
        import pytesseract
        OCR_AVAILABLE = True
        print("OCR support available")
    except ImportError:
        OCR_AVAILABLE = False
        print("OCR support not available - install pytesseract for image-based PDFs")
    
    print("All required packages imported successfully!")
except ImportError as e:
    print(f"Import error: {e}")
    print("Please install requirements: pip install -r requirements.txt")
    sys.exit(1)

# Configure logging for Railway (no file logging)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Production logging optimization
if os.environ.get('RAILWAY_ENVIRONMENT'):
    logger.setLevel(logging.WARNING)  # Reduce log noise in production

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
ALLOWED_EXTENSIONS = {'pdf'}
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

# Global API key storage
GEMINI_API_KEY = None

# Try to load local config for testing
try:
    from local_config import GEMINI_API_KEY as LOCAL_API_KEY
    if LOCAL_API_KEY:
        GEMINI_API_KEY = LOCAL_API_KEY
        print("✅ Local API key loaded for testing")
except ImportError:
    print("ℹ️ No local config found - API key will be set via web interface")

class EnhancedWatermarkRemover:
    """Advanced PDF watermark removal with AI-powered training and image detection"""
    
    def __init__(self):
        logger.info("Enhanced watermark remover initialized with AI-powered training and image detection")
        
        # Base patterns for common watermarks (fallback)
        self.base_patterns = [
            r'CONFIDENTIAL', r'DRAFT', r'COPYRIGHT', r'PROPRIETARY',
            r'RESTRICTED', r'PRIVATE', r'CLASSIFIED', r'TOP SECRET'
        ]
        
        # Trainable watermark patterns (learned from user documents)
        self.trained_patterns = []
        self.trained_watermarks = []
        
        # Auto-detection parameters
        self.watermark_threshold = 0.7  # Higher threshold - more conservative detection
        self.frequency_threshold = 5   # Higher frequency threshold - less aggressive
        self.position_weight = 0.3      # Weight for position-based detection
        
        # Training parameters
        self.min_training_samples = 3   # Minimum samples needed for training
        self.learning_rate = 0.1        # How much to adjust patterns based on feedback
        
        # AI Training parameters
        self.use_ai_training = True     # Enable AI-powered training
        self.ai_providers = ['gemini', 'openai']  # Supported AI providers
        self.ai_api_keys = {}  # Will be set via web interface - no hardcoded keys
        
        # Use global API key if available (for testing)
        global GEMINI_API_KEY
        if GEMINI_API_KEY:
            self.ai_api_keys['gemini'] = GEMINI_API_KEY
            logger.info("API key loaded from local config for testing")
        
    def set_api_key(self, provider: str, api_key: str):
        """Set API key for a specific provider"""
        if provider in self.ai_providers:
            self.ai_api_keys[provider] = api_key
            logger.info(f"API key set for {provider}")
            return True
        return False
        
        # Image detection parameters
        self.image_detection_enabled = True
        self.min_image_size = 100       # Minimum image size to detect (pixels)
        self.image_quality = 0.8        # Image quality for extraction
        
        # Load any existing trained patterns
        self._load_trained_patterns()
    
    def _ai_enhance_patterns(self, text: str, images: List, watermarks_to_remove: List, content_to_preserve: List) -> Dict:
        """Use AI to enhance watermark detection patterns"""
        try:
            enhanced_patterns = {'remove': [], 'preserve': []}
            
            # Try Gemini first (free tier available)
            if self._has_gemini_api():
                enhanced_patterns = self._gemini_enhance_patterns(text, images, watermarks_to_remove, content_to_preserve)
            # Fallback to OpenAI if available
            elif self._has_openai_api():
                enhanced_patterns = self._openai_enhance_patterns(text, images, watermarks_to_remove, content_to_preserve)
            else:
                logger.info("No AI API keys available, using basic training only")
                return enhanced_patterns
            
            return enhanced_patterns
            
        except Exception as e:
            logger.warning(f"AI enhancement failed: {e}")
            return {'remove': [], 'preserve': []}
    
    def _gemini_enhance_patterns(self, text: str, images: List, watermarks_to_remove: List, content_to_preserve: List) -> Dict:
        """Use Google Gemini to enhance watermark detection patterns"""
        try:
            import google.generativeai as genai
            
            # Configure Gemini
            genai.configure(api_key=self.ai_api_keys.get('gemini'))
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            # Create prompt for watermark analysis
            prompt = f"""
            Analyze this document content and identify additional watermarks that should be removed.
            
            Current watermarks to remove: {watermarks_to_remove}
            Content to preserve: {content_to_preserve}
            
            Document text (first 1000 chars): {text[:1000]}...
            
            Based on the content, identify:
            1. Additional watermark patterns (text that appears repeatedly, headers, footers, etc.)
            2. Content that should definitely be preserved (main text, questions, answers)
            3. Suggest improved patterns for existing watermarks
            
            Return as JSON:
            {{
                "remove": ["pattern1", "pattern2"],
                "preserve": ["content1", "content2"],
                "improvements": ["suggestion1", "suggestion2"]
            }}
            """
            
            response = model.generate_content(prompt)
            
            # Parse AI response
            try:
                import json
                result = json.loads(response.text)
                return {
                    'remove': result.get('remove', []),
                    'preserve': result.get('preserve', [])
                }
            except:
                # Fallback parsing
                return self._parse_ai_response_fallback(response.text)
                
        except Exception as e:
            logger.warning(f"Gemini enhancement failed: {e}")
            return {'remove': [], 'preserve': []}
    
    def _openai_enhance_patterns(self, text: str, images: List, watermarks_to_remove: List, content_to_preserve: List) -> Dict:
        """Use OpenAI ChatGPT to enhance watermark detection patterns"""
        try:
            import openai
            
            # Configure OpenAI
            openai.api_key = self.ai_api_keys.get('openai')
            
            # Create prompt for watermark analysis
            prompt = f"""
            Analyze this document content and identify additional watermarks that should be removed.
            
            Current watermarks to remove: {watermarks_to_remove}
            Content to preserve: {content_to_preserve}
            
            Document text (first 1000 chars): {text[:1000]}...
            
            Based on the content, identify:
            1. Additional watermark patterns (text that appears repeatedly, headers, footers, etc.)
            2. Content that should definitely be preserved (main text, questions, answers)
            3. Suggest improved patterns for existing watermarks
            
            Return as JSON:
            {{
                "remove": ["pattern1", "pattern2"],
                "preserve": ["content1", "content2"],
                "improvements": ["suggestion1", "suggestion2"]
            }}
            """
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an expert in document analysis and watermark detection."},
                    {"role": "user", "content": prompt}
                ]
            )
            
            ai_response = response.choices[0].message.content
            
            # Parse AI response
            try:
                import json
                result = json.loads(ai_response)
                return {
                    'remove': result.get('remove', []),
                    'preserve': result.get('preserve', [])
                }
            except:
                # Fallback parsing
                return self._parse_ai_response_fallback(ai_response)
                
        except Exception as e:
            logger.warning(f"OpenAI enhancement failed: {e}")
            return {'remove': [], 'preserve': []}
    
    def _parse_ai_response_fallback(self, ai_response: str) -> Dict:
        """Fallback parsing for AI responses that aren't valid JSON"""
        try:
            enhanced_patterns = {'remove': [], 'preserve': []}
            
            # Simple pattern extraction from text
            lines = ai_response.split('\n')
            for line in lines:
                line = line.strip().lower()
                if 'remove' in line or 'watermark' in line:
                    # Extract potential patterns
                    words = re.findall(r'\b\w+\b', line)
                    enhanced_patterns['remove'].extend(words[:3])  # Take first 3 words
                elif 'preserve' in line or 'keep' in line:
                    words = re.findall(r'\b\w+\b', line)
                    enhanced_patterns['preserve'].extend(words[:3])
            
            return enhanced_patterns
            
        except Exception as e:
            logger.warning(f"Fallback parsing failed: {e}")
            return {'remove': [], 'preserve': []}
    
    def _has_gemini_api(self) -> bool:
        """Check if Gemini API key is available"""
        return 'gemini' in self.ai_api_keys and self.ai_api_keys['gemini']
    
    def _has_openai_api(self) -> bool:
        """Check if OpenAI API key is available"""
        return 'openai' in self.ai_api_keys and self.ai_api_keys['openai']
    
    def _ai_enhance_watermark_removal(self, text: str) -> str:
        """Use Gemini AI to enhance watermark removal for SAT documents"""
        try:
            if not self._has_gemini_api():
                return None
            
            logger.info("Using Gemini AI to enhance watermark removal...")
            
            # Create prompt for Gemini
            prompt = f"""
            Analyze this SAT document text and identify watermarks to remove while preserving important content.
            
            Document text:
            {text[:2000]}...
            
            Instructions:
            1. Identify watermarks (headers, footers, page numbers, company names, etc.)
            2. Identify important content to preserve (reading passages, questions, multiple choice options)
            3. Return cleaned text with watermarks removed
            
            Return the cleaned text directly, no JSON formatting needed.
            """
            
            # Use Gemini API
            import google.generativeai as genai
            genai.configure(api_key=self.ai_api_keys['gemini'])
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            response = model.generate_content(prompt)
            cleaned_text = response.text
            
            if cleaned_text and len(cleaned_text) > 100:
                logger.info("Gemini AI successfully enhanced watermark removal")
                return cleaned_text
            else:
                logger.warning("Gemini AI response too short, using fallback")
                return None
                
        except Exception as e:
            logger.warning(f"Gemini AI enhancement failed: {e}")
            return None
    
    def set_ai_api_key(self, provider: str, api_key: str):
        """Set AI API key for a specific provider"""
        if provider in self.ai_providers:
            self.ai_api_keys[provider] = api_key
            logger.info(f"API key set for {provider}")
        else:
            logger.warning(f"Unsupported AI provider: {provider}")
    
    def _extract_images_from_pdf(self, pdf_path: str) -> List[Dict]:
        """Extract images from PDF with metadata"""
        try:
            images = []
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get image list for this page
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    try:
                        # Get image data
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Get image metadata
                        img_rect = page.get_image_bbox(img)
                        img_width = img_rect.width
                        img_height = img_rect.height
                        
                        # Only process images above minimum size
                        if img_width >= self.min_image_size and img_height >= self.min_image_size:
                            image_info = {
                                'page': page_num + 1,
                                'index': img_index,
                                'width': img_width,
                                'height': img_height,
                                'bbox': img_rect,
                                'data': image_bytes,
                                'format': base_image["ext"],
                                'size': len(image_bytes)
                            }
                            images.append(image_info)
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract image {img_index} from page {page_num}: {e}")
                        continue
            
            doc.close()
            logger.info(f"Extracted {len(images)} images from PDF")
            return images
            
        except Exception as e:
            logger.error(f"Image extraction failed: {e}")
            return []
    
    def train_on_document(self, document_path: str, user_feedback: Dict) -> bool:
        """Train the watermark detector on a specific document with AI-powered enhancement"""
        try:
            logger.info(f"Training watermark detector on: {document_path}")
            
            # Extract text and images from document
            text = self._extract_text_enhanced(document_path)
            images = self._extract_images_from_pdf(document_path)
            
            # Get user feedback about what should/shouldn't be removed
            watermarks_to_remove = user_feedback.get('remove', [])
            content_to_preserve = user_feedback.get('preserve', [])
            
            # AI-powered pattern enhancement
            if self.use_ai_training:
                enhanced_patterns = self._ai_enhance_patterns(
                    text, images, watermarks_to_remove, content_to_preserve
                )
                watermarks_to_remove.extend(enhanced_patterns.get('remove', []))
                content_to_preserve.extend(enhanced_patterns.get('preserve', []))
                logger.info(f"AI enhanced patterns: {len(enhanced_patterns.get('remove', []))} removal, {len(enhanced_patterns.get('preserve', []))} preservation")
            
            # Learn new patterns
            new_patterns = self._learn_from_feedback(text, watermarks_to_remove, content_to_preserve)
            
            # Update trained patterns
            self.trained_patterns.extend(new_patterns)
            
            # Save trained patterns
            self._save_trained_patterns()
            
            logger.info(f"Training completed. New patterns learned: {len(new_patterns)}")
            return True
            
        except Exception as e:
            logger.error(f"Training failed: {e}")
            return False
    
    def _learn_from_feedback(self, text: str, watermarks_to_remove: List[str], content_to_preserve: List[str]) -> List[str]:
        """Learn new patterns from user feedback"""
        new_patterns = []
        
        # Learn patterns for watermarks to remove
        for watermark in watermarks_to_remove:
            pattern = self._create_pattern_from_text(watermark)
            if pattern:
                new_patterns.append(pattern)
                logger.info(f"Learned removal pattern: {pattern}")
        
        # Learn patterns for content to preserve (negative learning)
        for content in content_to_preserve:
            # Create a pattern that should NOT match
            pattern = self._create_preservation_pattern(content)
            if pattern:
                new_patterns.append(pattern)
                logger.info(f"Learned preservation pattern: {pattern}")
        
        return new_patterns
    
    def _create_pattern_from_text(self, text: str) -> str:
        """Create a regex pattern from text for watermark detection"""
        if not text or len(text) < 2:
            return None
        
        # Escape special characters
        escaped_text = re.escape(text)
        
        # Make it flexible for variations
        pattern = f"\\b{escaped_text}\\b"
        
        return pattern
    
    def _create_preservation_pattern(self, text: str) -> str:
        """Create a pattern that should preserve content"""
        if not text or len(text) < 2:
            return None
        
        # This is a negative pattern - we'll handle it specially
        escaped_text = re.escape(text)
        pattern = f"PRESERVE:{escaped_text}"
        
        return pattern
    
    def _load_trained_patterns(self):
        """Load trained patterns from storage"""
        try:
            # For now, load from memory. In production, this would be from a database/file
            # This is where you'd implement persistent storage
            pass
        except Exception as e:
            logger.warning(f"Failed to load trained patterns: {e}")
    
    def _save_trained_patterns(self):
        """Save trained patterns to storage"""
        try:
            # For now, save to memory. In production, this would be to a database/file
            # This is where you'd implement persistent storage
            pass
        except Exception as e:
            logger.warning(f"Failed to save trained patterns: {e}")
    
    def learn_from_approval(self, document_path: str, user_feedback: Dict) -> bool:
        """Learn from user approval of successful output"""
        try:
            logger.info(f"Learning from approval for: {document_path}")
            
            # Extract text and images from document
            text = self._extract_text_enhanced(document_path)
            images = self._extract_images_from_pdf(document_path)
            
            # Get user feedback about what worked well
            approved_content = user_feedback.get('preserve', [])
            user_comment = user_feedback.get('comment', '')
            
            # AI-powered pattern enhancement for successful patterns
            if self.use_ai_training:
                enhanced_patterns = self._ai_enhance_successful_patterns(
                    text, images, approved_content, user_comment
                )
                approved_content.extend(enhanced_patterns.get('preserve', []))
                logger.info(f"AI enhanced successful patterns: {len(enhanced_patterns.get('preserve', []))} preservation patterns")
            
            # Learn successful patterns
            successful_patterns = self._learn_from_success(text, approved_content, user_comment)
            
            # Update trained patterns with successful ones
            self.trained_patterns.extend(successful_patterns)
            
            # Save trained patterns
            self._save_trained_patterns()
            
            logger.info(f"Learning from approval completed. New successful patterns learned: {len(successful_patterns)}")
            return True
            
        except Exception as e:
            logger.error(f"Learning from approval failed: {e}")
            return False
    
    def _learn_from_success(self, text: str, approved_content: List[str], user_comment: str) -> List[str]:
        """Learn patterns from successful output"""
        successful_patterns = []
        
        # Learn patterns for content that was approved
        for content in approved_content:
            pattern = self._create_success_pattern(content)
            if pattern:
                successful_patterns.append(pattern)
                logger.info(f"Learned success pattern: {pattern}")
        
        # Learn from user comments about what worked
        if user_comment:
            comment_patterns = self._extract_patterns_from_comment(user_comment)
            successful_patterns.extend(comment_patterns)
            logger.info(f"Learned patterns from user comment: {len(comment_patterns)}")
        
        return successful_patterns
    
    def _create_success_pattern(self, content: str) -> str:
        """Create a pattern from successful content"""
        if not content or len(content) < 2:
            return None
        
        # Escape special characters
        escaped_content = re.escape(content)
        
        # Make it flexible for variations
        pattern = f"SUCCESS:{escaped_content}"
        
        return pattern
    
    def _extract_patterns_from_comment(self, comment: str) -> List[str]:
        """Extract patterns from user comment"""
        patterns = []
        
        # Look for specific mentions of what worked well
        if 'good' in comment.lower() or 'correct' in comment.lower():
            # Extract context around positive words
            words = comment.split()
            for i, word in enumerate(words):
                if word.lower() in ['good', 'correct', 'perfect', 'excellent']:
                    # Get surrounding context
                    start = max(0, i-2)
                    end = min(len(words), i+3)
                    context = ' '.join(words[start:end])
                    pattern = f"SUCCESS_CONTEXT:{re.escape(context)}"
                    patterns.append(pattern)
        
        return patterns
    
    def _ai_enhance_successful_patterns(self, text: str, images: List[Dict], approved_content: List[str], user_comment: str) -> Dict:
        """Use AI to enhance successful patterns"""
        try:
            if not self.ai_api_keys:
                return {'preserve': []}
            
            # Prepare AI prompt for successful pattern enhancement
            prompt = f"""
            Analyze this successful SAT document processing and identify patterns that worked well:
            
            Text: {text[:1000]}...
            Approved Content: {approved_content}
            User Comment: {user_comment}
            
            Identify patterns that should be preserved in future processing:
            1. Text patterns that were correctly identified as content
            2. Question structures that were properly detected
            3. Multiple choice formats that were correctly parsed
            4. Reading passage indicators that worked well
            
            Return as JSON with 'preserve' array of patterns to keep.
            """
            
            # Use Gemini AI for enhancement
            for api_key in self.ai_api_keys.values():
                try:
                    import google.generativeai as genai
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel('gemini-2.5-flash')
                    
                    response = model.generate_content(prompt)
                    result = json.loads(response.text)
                    
                    logger.info(f"AI enhanced successful patterns: {len(result.get('preserve', []))} patterns")
                    return result
                    
                except Exception as e:
                    logger.warning(f"AI enhancement failed with key: {e}")
                    continue
            
            return {'preserve': []}
            
        except Exception as e:
            logger.warning(f"AI enhancement of successful patterns failed: {e}")
            return {'preserve': []}
    
    def remove_watermarks_from_pdf(self, input_path: str, output_path: str) -> bool:
        """
        Enhanced watermark removal using multiple detection methods
        Based on techniques from lxulxu/WatermarkRemover and marcbelmont/cnn-watermark-removal
        """
        try:
            logger.info(f"Starting enhanced watermark removal from {input_path}")
            
            # Open the PDF
            doc = fitz.open(input_path)
            
            # Store cleaned text for each page
            cleaned_pages = []
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                logger.info(f"Processing page {page_num + 1}")
                
                # Get page text with enhanced extraction
                text = self._extract_text_enhanced(page)
                
                # Enhanced watermark detection and removal
                cleaned_text = self._remove_watermarks_enhanced(text)
                
                # Store cleaned text for this page
                cleaned_pages.append({
                    'page_num': page_num + 1,
                    'original_text': text,
                    'cleaned_text': cleaned_text
                })
            
            # Close the original PDF
            doc.close()
            
            # Create a new PDF with cleaned content
            new_doc = fitz.open()
            
            for page_data in cleaned_pages:
                # Create new page
                new_page = new_doc.new_page()
                
                # Insert cleaned text with better formatting
                new_page.insert_text((50, 50), page_data['cleaned_text'], fontsize=11)
            
            # Save the cleaned PDF
            new_doc.save(output_path)
            new_doc.close()
            
            logger.info(f"Enhanced watermark removal completed: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Enhanced watermark removal failed: {str(e)}")
            logger.error(traceback.format_exc())
            return False
    
    def _extract_text_enhanced(self, page) -> str:
        """Enhanced text extraction with better OCR-like processing"""
        try:
            # Get text with better extraction settings
            text = page.get_text("text")
            
            # Additional text blocks for better coverage
            blocks = page.get_text("dict")
            additional_text = ""
            
            for block in blocks["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if span["text"].strip():
                                additional_text += span["text"] + " "
            
            # Combine and clean
            combined_text = text + " " + additional_text
            return combined_text.strip()
            
        except Exception as e:
            logger.warning(f"Enhanced text extraction failed, using basic: {e}")
            return page.get_text()
    
    def extract_text_from_pdf(self, pdf_path: str, max_pages: int = None) -> str:
        """Extract text from PDF file with enhanced processing and OCR support"""
        try:
            logger.info(f"Extracting text from PDF: {pdf_path}")
            
            # Open the PDF
            doc = fitz.open(pdf_path)
            
            # Determine how many pages to process
            if max_pages is None:
                max_pages = len(doc)
            else:
                max_pages = min(max_pages, len(doc))
            
            # Extract text from each page
            all_text = []
            for page_num in range(max_pages):
                page = doc[page_num]
                page_text = self._extract_text_enhanced(page)
                
                # If no text found, try OCR for image-based PDFs
                if not page_text.strip():
                    logger.info(f"No text found on page {page_num + 1}, trying OCR...")
                    page_text = self._extract_text_with_ocr(page)
                
                all_text.append(page_text)
            
            # Close the document
            doc.close()
            
            # Combine all text
            combined_text = "\n".join(all_text)
            logger.info(f"Extracted {len(combined_text)} characters from {max_pages} pages")
            
            return combined_text
            
        except Exception as e:
            logger.error(f"Text extraction from PDF failed: {e}")
            return ""
    
    def _extract_text_with_ocr(self, page) -> str:
        """Extract text from page using OCR for image-based PDFs"""
        try:
            # Check if OCR is available
            if not OCR_AVAILABLE:
                logger.warning("OCR not available - skipping image-based text extraction")
                return ""
            
            # Convert page to image
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Convert to PIL Image
            image = Image.open(io.BytesIO(img_data))
            
            # Use OCR to extract text
            text = pytesseract.image_to_string(image, config='--psm 6')
            
            # Clean up the text
            text = self._clean_ocr_text(text)
            
            logger.info(f"OCR extracted {len(text)} characters")
            return text
            
        except Exception as e:
            logger.warning(f"OCR extraction failed: {e}")
            return ""
    
    def _clean_ocr_text(self, text: str) -> str:
        """Clean up OCR text to improve quality"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR errors
        text = re.sub(r'[|]', 'I', text)  # Fix pipe to I
        text = re.sub(r'[0]', 'O', text)  # Fix 0 to O in words
        text = re.sub(r'[1]', 'l', text)  # Fix 1 to l in words
        
        # Remove page numbers and headers/footers
        text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
        
        return text.strip()
    
    def _remove_watermarks_enhanced(self, text: str) -> str:
        """Advanced watermark removal with automatic detection and AI enhancement"""
        try:
            logger.info("Starting automatic watermark detection and removal with AI")
            
            # Step 1: Use AI to enhance watermark detection if available
            if self._has_gemini_api():
                ai_enhanced_text = self._ai_enhance_watermark_removal(text)
                if ai_enhanced_text:
                    logger.info("AI-enhanced watermark removal completed")
                    return ai_enhanced_text
            
            # Step 2: Analyze text structure and extract potential watermarks
            potential_watermarks = self._detect_watermarks_automatically(text)
            logger.info(f"Detected {len(potential_watermarks)} potential watermarks")
            
            # Step 3: Remove detected watermarks
            cleaned_text = self._remove_detected_watermarks(text, potential_watermarks)
            
            # Step 4: Clean up formatting
            cleaned_text = self._clean_text_formatting(cleaned_text)
            
            return cleaned_text
            
        except Exception as e:
            logger.warning(f"Auto watermark detection failed, using fallback: {e}")
            return self._remove_watermarks_fallback(text)
    
    def _detect_watermarks_automatically(self, text: str) -> List[Dict]:
        """Automatically detect watermarks using multiple algorithms"""
        watermarks = []
        
        # Method 1: Frequency-based detection
        freq_watermarks = self._detect_frequency_watermarks(text)
        watermarks.extend(freq_watermarks)
        
        # Method 2: Position-based detection
        pos_watermarks = self._detect_position_watermarks(text)
        watermarks.extend(pos_watermarks)
        
        # Method 3: Pattern-based detection
        pattern_watermarks = self._detect_pattern_watermarks(text)
        watermarks.extend(pattern_watermarks)
        
        # Method 4: Context-based detection
        context_watermarks = self._detect_context_watermarks(text)
        watermarks.extend(context_watermarks)
        
        # Remove duplicates and return
        unique_watermarks = self._deduplicate_watermarks(watermarks)
        return unique_watermarks
    
    def _detect_frequency_watermarks(self, text: str) -> List[Dict]:
        """Detect watermarks based on word frequency analysis"""
        words = re.findall(r'\b\w+\b', text.lower())
        word_freq = {}
        
        # Count frequencies
        for word in words:
            if len(word) > 2:  # Skip very short words
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Find words that appear too frequently (likely watermarks)
        watermarks = []
        total_words = len(words)
        
        for word, freq in word_freq.items():
            if freq >= self.frequency_threshold and freq / total_words > 0.05:
                watermarks.append({
                    'type': 'frequency',
                    'text': word,
                    'confidence': min(freq / total_words, 1.0),
                    'pattern': r'\b' + re.escape(word) + r'\b'
                })
        
        return watermarks
    
    def _detect_position_watermarks(self, text: str) -> List[Dict]:
        """Detect watermarks based on position (headers/footers)"""
        lines = text.split('\n')
        watermarks = []
        
        # Check first and last few lines
        header_lines = lines[:3]
        footer_lines = lines[-3:] if len(lines) > 3 else []
        
        for line in header_lines + footer_lines:
            line = line.strip()
            if line and len(line) < 50:  # Short lines in headers/footers
                # Check if line contains watermark-like content
                if self._is_watermark_like(line):
                    watermarks.append({
                        'type': 'position',
                        'text': line,
                        'confidence': 0.8,
                        'pattern': re.escape(line)
                    })
        
        return watermarks
    
    def _detect_pattern_watermarks(self, text: str) -> List[Dict]:
        """Detect watermarks using pattern matching"""
        watermarks = []
        
        # Check against base patterns
        for pattern in self.base_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                watermarks.append({
                    'type': 'pattern',
                    'text': match,
                    'confidence': 0.9,
                    'pattern': r'\b' + re.escape(match) + r'\b'
                })
        
        # Detect repeated phrases
        repeated_phrases = self._find_repeated_phrases(text)
        watermarks.extend(repeated_phrases)
        
        return watermarks
    
    def _detect_context_watermarks(self, text: str) -> List[Dict]:
        """Detect watermarks based on context and surrounding text"""
        watermarks = []
        
        # Look for copyright notices, page numbers, etc.
        context_patterns = [
            (r'©\s*\d{4}.*?(?:all rights reserved|inc\.|llc\.)', 0.9),
            (r'page\s+\d+\s+of\s+\d+', 0.8),
            (r'generated\s+on.*?\d{4}', 0.8),
            (r'last\s+modified.*?\d{4}', 0.8),
            (r'confidential.*?document', 0.9),
            (r'proprietary.*?information', 0.9)
        ]
        
        for pattern, confidence in context_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                watermarks.append({
                    'type': 'context',
                    'text': match,
                    'confidence': confidence,
                    'pattern': pattern
                })
        
        return watermarks
    
    def _is_watermark_like(self, text: str) -> bool:
        """Check if text looks like a watermark"""
        # Short text
        if len(text) > 50:
            return False
        
        # Contains common watermark words
        watermark_indicators = [
            'confidential', 'draft', 'copyright', 'proprietary',
            'restricted', 'private', 'classified', 'internal',
            'watermark', 'stamp', 'logo', 'brand'
        ]
        
        text_lower = text.lower()
        for indicator in watermark_indicators:
            if indicator in text_lower:
                return True
        
        # All caps (common for watermarks) - but be more conservative
        if text.isupper() and len(text) > 5 and len(text) < 20:
            return True
        
        # Contains special characters or formatting
        if re.search(r'[©®™]', text):
            return True
        
        # Don't mark short text as watermark unless it's clearly a watermark
        if len(text) < 10:
            return False
        
        return False
    
    def _find_repeated_phrases(self, text: str) -> List[Dict]:
        """Find phrases that repeat frequently (likely watermarks)"""
        watermarks = []
        
        # Look for 2-4 word phrases that repeat
        words = re.findall(r'\b\w+\b', text.lower())
        
        for phrase_length in range(2, 5):
            phrases = {}
            for i in range(len(words) - phrase_length + 1):
                phrase = ' '.join(words[i:i + phrase_length])
                if len(phrase) > 10:  # Skip very short phrases
                    phrases[phrase] = phrases.get(phrase, 0) + 1
            
            # Find frequently repeated phrases
            for phrase, count in phrases.items():
                if count >= 3:  # Appears at least 3 times
                    watermarks.append({
                        'type': 'repeated_phrase',
                        'text': phrase,
                        'confidence': min(count / 10, 1.0),
                        'pattern': re.escape(phrase)
                    })
        
        return watermarks
    
    def _deduplicate_watermarks(self, watermarks: List[Dict]) -> List[Dict]:
        """Remove duplicate watermarks and merge similar ones"""
        unique_watermarks = []
        seen_texts = set()
        
        # Sort by confidence (highest first)
        watermarks.sort(key=lambda x: x['confidence'], reverse=True)
        
        for watermark in watermarks:
            text_lower = watermark['text'].lower()
            if text_lower not in seen_texts:
                unique_watermarks.append(watermark)
                seen_texts.add(text_lower)
        
        return unique_watermarks
    
    def _remove_detected_watermarks(self, text: str, watermarks: List[Dict]) -> str:
        """Remove detected watermarks from text"""
        cleaned_text = text
        
        for watermark in watermarks:
            if watermark['confidence'] >= self.watermark_threshold:
                pattern = watermark['pattern']
                cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE)
                logger.info(f"Removed watermark: '{watermark['text']}' (confidence: {watermark['confidence']:.2f})")
        
        return cleaned_text
    
    def _remove_watermarks_fallback(self, text: str) -> str:
        """Fallback watermark removal using basic patterns"""
        cleaned_text = text
        
        # Use base patterns as fallback
        for pattern in self.base_patterns:
            pattern_with_boundaries = r'\b' + pattern + r'\b'
            cleaned_text = re.sub(pattern_with_boundaries, '', cleaned_text, flags=re.IGNORECASE)
        
        return re.sub(r'\s+', ' ', cleaned_text).strip()
    
    def _clean_text_formatting(self, text: str) -> str:
        """Clean up text formatting and spacing"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR issues
        replacements = [
            ('||', 'll'), ('|/', 'll'), ('0O', '0'), ('O0', '0'),
            ('1l', 'll'), ('l1', 'll'), ('5S', 'S'), ('S5', 'S'),
            ('rn', 'm'), ('cl', 'd'), ('vv', 'w'), ('nn', 'm')
        ]
        
        for old, new in replacements:
            text = text.replace(old, new)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    


class SATDocumentProcessor:
    """Specialized SAT document processor with format detection and preservation"""
    
    def __init__(self):
        self.watermark_remover = EnhancedWatermarkRemover()
        self.sat_patterns = self._initialize_sat_patterns()
        logger.info("SAT document processor initialized")
    
    def _initialize_sat_patterns(self):
        """Initialize enhanced SAT-specific detection patterns"""
        return {
            'section_headers': [
                r'^SAT\s+Practice\s+Test',  # "SAT Practice Test"
                r'^(?:Section|Part)\s+\d+[:\s]*',  # "Section 1:" or "Part 1:"
                r'^(?:Reading|Writing|Math|Language)\s+(?:Test|Section)[:\s]*',  # "Reading Test:" or "Math Section:"
                r'^(?:Evidence-Based|Critical\s+Reading|Mathematics)[:\s]*',  # "Evidence-Based Reading:" or "Mathematics:"
                r'^(?:English|Math|Reading|Writing)\s+Section',  # "English Section" or "Math Section"
            ],
            'reading_passage': [
                # Digital SAT format patterns (shorter passages, 25-150 words)
                r'^Reading\s+Passage\s*\d+[:\s]*$',  # "Reading Passage 1:" at start of line
                r'^Questions?\s*\d+[-\s]*\d*\s*are\s+based\s+on\s+the\s+following\s+passage$',  # "Questions 1-10 are based on the following passage"
                r'^The\s+following\s+passage\s+is\s+adapted\s+from',  # "The following passage is adapted from..."
                r'^Read\s+the\s+following\s+passage',  # "Read the following passage..."
                # Enhanced digital SAT patterns based on research
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 6+ word sentences (typical passage start)
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 5+ word sentences
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 4+ word sentences
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+',  # 3+ word sentences
                # Common passage starters for digital SAT
                r'^The\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # "The unique subak water..."
                r'^Until\s+\d{4}',  # "Until 1917, there was..."
                r'^[A-Z][a-z]+\s+[a-z]+\s+\([^)]+\)',  # "The mihrab (or niche)..."
                r'^[A-Z][a-z]+\s+[a-z]+-[a-z]+',  # "The Egyptian plover-a bird..."
                # Digital SAT reading passage indicators (longer sentences)
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 7+ word sentences
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 8+ word sentences
                # Flexible passage detection for digital SAT
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 9+ word sentences
                # Additional digital SAT patterns
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 10+ word sentences
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 11+ word sentences
            ],
            'question_start': [
                # Digital SAT question patterns
                r'^\s*\d+\.\s+',  # "1. " at start of line
                r'^\s*Question\s+\d+[:\s]*',  # "Question 1:" or "Question 1"
                r'^\s*Q\s*\d+[:\s]*',  # "Q1:" or "Q1"
                r'^\s*\d+\s*$',  # "31" (standalone number)
                r'^\s*\d+\s+[A-Z]',  # "31 A" (number followed by letter)
                r'^\s*[A-Z]\s*$',  # "A" (standalone letter)
                # Enhanced digital SAT question detection
                r'^\s*\d+\s*[A-Z]',  # "1 A" (number followed by letter)
                r'^\s*\d+\s*[a-z]',  # "1 a" (number followed by lowercase letter)
                # Digital SAT specific patterns
                r'^\s*\d+\s*[A-Z][a-z]',  # "1 Which" (number followed by question word)
                r'^\s*\d+\s*[A-Z][a-z]+\s+[a-z]+',  # "1 Which choice" (number followed by question phrase)
                r'^\s*\d+\s*[A-Z][a-z]+\s+[a-z]+\s+[a-z]+',  # "1 Which choice best" (number followed by longer question phrase)
            ],
            'multiple_choice': [
                r'^\s*[A-D]\)\s+',  # "A) ", "B) ", "C) ", "D) "
                r'^\s*[A-D]\.\s+',  # "A. ", "B. ", "C. ", "D. "
                r'^\s*[A-D]\s+',    # "A ", "B ", "C ", "D "
                r'^\s*[A-D]\s*$',   # "A", "B", "C", "D" (standalone)
                r'^\s*[A-D]\s*[A-Z]',  # "A A+", "B B+" (with additional text)
                r'^\s*[A-D]\s*[-+]\s*[A-Z]',  # "A- A+", "B- B+" (with dash)
                r'^\s*[A-D]\s*[A-Z]+\s*$',  # "A A+", "B B+" (standalone with letters)
                # Enhanced multiple choice detection
                r'^\s*[A-D]\s*[a-z]',  # "A a", "B b" (with lowercase)
                r'^\s*[A-D]\s*[0-9]',  # "A 1", "B 2" (with numbers)
                r'^\s*[A-D]\s*[^\s]',  # "A +", "B -" (with symbols)
            ],
            'math_indicators': [
                r'\b(?:function|equation|graph|slope|intercept|variable|solve|calculate|area|perimeter|volume|angle|triangle|rectangle|circle)\b',
                r'[=+\-*/(){}[\]^]',  # Mathematical symbols
                r'\b\d+\s*(?:times|multiplied by|divided by|plus|minus)\b',  # Word problems
                r'\b(?:length|width|height|radius|diameter|base|height)\b',  # Geometry terms
                # Enhanced math detection
                r'\b(?:algebra|geometry|trigonometry|calculus|statistics|probability)\b',
                r'\b(?:x|y|z|a|b|c)\s*[=+\-*/]',  # Variables in equations
                r'\b(?:fraction|decimal|percentage|ratio|proportion)\b',
                r'\b(?:quadratic|linear|exponential|logarithmic)\b',
            ],
            'reading_indicators': [
                r'\b(?:passage|author|text|paragraph|line|suggests|implies|indicates|according to)\b',
                r'\b(?:main idea|central theme|primary purpose|best describes)\b',
                r'\b(?:inference|conclusion|implication|evidence)\b',
                r'\b(?:vocabulary|word|phrase|meaning|definition)\b',
                # Enhanced reading detection
                r'\b(?:tone|mood|style|rhetoric|argument|persuasion)\b',
                r'\b(?:compare|contrast|analyze|evaluate|interpret)\b',
                r'\b(?:support|refute|challenge|question)\b',
            ],
            'writing_indicators': [
                r'\b(?:grammar|sentence|paragraph|transition|conclusion|introduction)\b',
                r'\b(?:subject|verb|pronoun|agreement|tense|parallel)\b',
                r'\b(?:no change|omit|delete|add|replace)\b',
                r'[A-D]\)\s*(?:No change|omit|delete)',  # Writing section options
                # Enhanced writing detection
                r'\b(?:punctuation|comma|semicolon|colon|apostrophe)\b',
                r'\b(?:conjunction|preposition|article|adjective|adverb)\b',
                r'\b(?:clause|phrase|sentence|paragraph|essay)\b',
            ],
            'image_indicators': [
                r'\[Image\]',  # "[Image]" placeholder
                r'\[Figure\s+\d+\]',  # "[Figure 1]" placeholder
                r'\[Graph\]',  # "[Graph]" placeholder
                r'\[Chart\]',  # "[Chart]" placeholder
                r'\[Diagram\]',  # "[Diagram]" placeholder
            ]
        }
    
    def _classify_question_type(self, question_text: str, context: str = "") -> Dict:
        """Classify SAT question type based on content analysis"""
        try:
            question_lower = question_text.lower()
            context_lower = context.lower()
            combined_text = f"{question_lower} {context_lower}"
            
            # Initialize classification
            classification = {
                'question_type': 'unknown',
                'section_type': 'unknown',
                'difficulty': 'medium',
                'topic': 'general',
                'format_requirements': []
            }
            
            # Math question detection
            math_score = 0
            for pattern in self.sat_patterns['math_indicators']:
                if re.search(pattern, combined_text, re.IGNORECASE):
                    math_score += 1
            
            # Reading question detection
            reading_score = 0
            for pattern in self.sat_patterns['reading_indicators']:
                if re.search(pattern, combined_text, re.IGNORECASE):
                    reading_score += 1
            
            # Writing question detection
            writing_score = 0
            for pattern in self.sat_patterns['writing_indicators']:
                if re.search(pattern, combined_text, re.IGNORECASE):
                    writing_score += 1
            
            # Determine question type based on scores
            if math_score > reading_score and math_score > writing_score:
                classification['section_type'] = 'math'
                classification['format_requirements'].append('preserve_math_notation')
                
                # Sub-classify math questions
                if re.search(r'\b(?:function|equation|graph)\b', combined_text):
                    classification['question_type'] = 'algebra'
                    classification['topic'] = 'functions_and_equations'
                elif re.search(r'\b(?:area|perimeter|volume|angle|triangle|rectangle|circle)\b', combined_text):
                    classification['question_type'] = 'geometry'
                    classification['topic'] = 'geometry'
                elif re.search(r'\b(?:data|chart|graph|statistics|probability)\b', combined_text):
                    classification['question_type'] = 'data_analysis'
                    classification['topic'] = 'data_analysis'
                else:
                    classification['question_type'] = 'problem_solving'
                    classification['topic'] = 'word_problems'
                    
            elif reading_score > writing_score:
                classification['section_type'] = 'reading'
                classification['format_requirements'].append('preserve_passage_structure')
                
                # Sub-classify reading questions
                if re.search(r'\b(?:main idea|central theme|primary purpose)\b', combined_text):
                    classification['question_type'] = 'main_idea'
                    classification['topic'] = 'comprehension'
                elif re.search(r'\b(?:according to|states|mentions)\b', combined_text):
                    classification['question_type'] = 'detail'
                    classification['topic'] = 'detail_retrieval'
                elif re.search(r'\b(?:suggests|implies|indicates|inference)\b', combined_text):
                    classification['question_type'] = 'inference'
                    classification['topic'] = 'inference'
                elif re.search(r'\b(?:vocabulary|word|phrase|meaning)\b', combined_text):
                    classification['question_type'] = 'vocabulary'
                    classification['topic'] = 'vocabulary_in_context'
                elif re.search(r'\b(?:evidence|support|best evidence)\b', combined_text):
                    classification['question_type'] = 'evidence'
                    classification['topic'] = 'command_of_evidence'
                else:
                    classification['question_type'] = 'comprehension'
                    classification['topic'] = 'general_comprehension'
                    
            elif writing_score > 0:
                classification['section_type'] = 'writing'
                classification['format_requirements'].append('maintain_line_numbers')
                
                # Sub-classify writing questions
                if re.search(r'\b(?:grammar|subject|verb|pronoun|agreement)\b', combined_text):
                    classification['question_type'] = 'grammar'
                    classification['topic'] = 'grammar_and_usage'
                elif re.search(r'\b(?:sentence|paragraph|transition|organization)\b', combined_text):
                    classification['question_type'] = 'expression'
                    classification['topic'] = 'expression_of_ideas'
                else:
                    classification['question_type'] = 'editing'
                    classification['topic'] = 'standard_english_conventions'
            
            # Determine difficulty based on question complexity
            if len(question_text) > 100 or re.search(r'\b(?:complex|advanced|sophisticated)\b', combined_text):
                classification['difficulty'] = 'hard'
            elif len(question_text) < 50:
                classification['difficulty'] = 'easy'
            
            logger.info(f"Question classified as: {classification['section_type']} - {classification['question_type']}")
            return classification
            
        except Exception as e:
            logger.warning(f"Question classification failed: {e}")
            return {
                'question_type': 'unknown',
                'section_type': 'unknown',
                'difficulty': 'medium',
                'topic': 'general',
                'format_requirements': []
            }
    
    def _ai_enhance_sat_structure_detection(self, text: str) -> Dict:
        """Use Gemini AI to enhance SAT structure detection"""
        try:
            if not self.watermark_remover._has_gemini_api():
                return None
            
            logger.info("Using Gemini AI to enhance SAT structure detection...")
            
            # Create prompt for Gemini
            prompt = f"""
            Analyze this SAT document text and identify the structure with exact formatting.
            
            Document text:
            {text[:3000]}...
            
            Instructions:
            1. Identify reading passages (long text blocks before questions)
            2. Identify questions (numbered items that ask something)
            3. Identify multiple choice options (A, B, C, D choices)
            4. Return the structure in this exact format:
            
            {{
                "reading_passages": [
                    {{
                        "text": "passage title or first line",
                        "content": ["line 1", "line 2", "line 3"]
                    }}
                ],
                "questions": [
                    {{
                        "text": "question text",
                        "content": ["question content"],
                        "choices": [
                            {{"text": "A) choice text"}},
                            {{"text": "B) choice text"}},
                            {{"text": "C) choice text"}},
                            {{"text": "D) choice text"}}
                        ]
                    }}
                ]
            }}
            
            Focus on the exact format above. Return valid JSON only.
            """
            
            # Use Gemini API
            import google.generativeai as genai
            genai.configure(api_key=self.watermark_remover.ai_api_keys['gemini'])
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            response = model.generate_content(prompt)
            ai_response = response.text
            
            # Parse AI response
            try:
                import json
                result = json.loads(ai_response)
                
                # Validate structure
                if 'reading_passages' in result and 'questions' in result:
                    logger.info("Gemini AI successfully enhanced SAT structure detection")
                    return result
                else:
                    logger.warning("Gemini AI response missing required fields")
                    return None
                    
            except json.JSONDecodeError as e:
                logger.warning(f"Failed to parse Gemini AI response as JSON: {e}")
                return None
                
        except Exception as e:
            logger.warning(f"Gemini AI enhancement failed: {e}")
            return None
    
    def process_sat_document(self, input_path: str) -> str:
        """Process SAT document with format detection, image preservation, and watermark removal"""
        try:
            logger.info(f"Processing SAT document: {input_path}")
            
            # Extract text with watermark removal
            text = self.watermark_remover.extract_text_from_pdf(input_path)
            logger.info(f"Extracted text length: {len(text)} characters")
            
            # Extract images for integration (disabled for now due to extraction issues)
            images = []  # self.watermark_remover._extract_images_from_pdf(input_path)
            logger.info(f"Extracted {len(images)} images from document")
            
            # Detect SAT structure with image integration
            sat_structure = self._detect_sat_structure_with_images(text, images)
            logger.info(f"Detected SAT structure with images: {sat_structure}")
            
            # Process and format for Word with images
            formatted_text = self._format_sat_for_word_with_images(text, sat_structure, images)
            
            return formatted_text
            
        except Exception as e:
            logger.error(f"SAT document processing failed: {e}")
            raise
    
    def process_sat_document_english(self, input_path: str) -> str:
        """Process SAT document specifically for English section (text processing focus)"""
        try:
            logger.info(f"Processing SAT document for English section: {input_path}")
            
            # Extract text with enhanced watermark removal for English
            text = self.watermark_remover.extract_text_from_pdf(input_path)
            logger.info(f"Extracted text length: {len(text)} characters")
            
            # Extract images for integration (disabled for now due to extraction issues)
            images = []  # self.watermark_remover._extract_images_from_pdf(input_path)
            logger.info(f"Extracted {len(images)} images from document")
            
            # Detect SAT structure with image integration
            sat_structure = self._detect_sat_structure_with_images(text, images)
            logger.info(f"Detected SAT structure with images: {sat_structure}")
            
            # Process and format for Word with English section focus
            formatted_text = self._format_sat_for_word_english(text, sat_structure, images)
            
            return formatted_text
            
        except Exception as e:
            logger.error(f"SAT document processing for English section failed: {e}")
            raise
    
    def process_sat_document_math(self, input_path: str) -> str:
        """Process SAT document specifically for Math section (LaTeX conversion focus)"""
        try:
            logger.info(f"Processing SAT document for Math section: {input_path}")
            
            # Extract text with enhanced watermark removal for Math
            text = self.watermark_remover.extract_text_from_pdf(input_path)
            logger.info(f"Extracted text length: {len(text)} characters")
            
            # Extract images for integration
            images = self.watermark_remover._extract_images_from_pdf(input_path)
            logger.info(f"Extracted {len(text)} images from document")
            
            # Detect SAT structure with image integration
            sat_structure = self._detect_sat_structure_with_images(text, images)
            logger.info(f"Detected SAT structure with images: {sat_structure}")
            
            # Process and format for Word with Math section focus (LaTeX conversion)
            formatted_text = self._format_sat_for_word_math(text, sat_structure, images)
            
            return formatted_text
            
        except Exception as e:
            logger.error(f"SAT document processing for Math section failed: {e}")
            raise
    
    def process_sat_document_unified(self, input_path: str) -> str:
        """Process SAT document with unified English/Math processing"""
        try:
            logger.info(f"Processing SAT document with unified processing: {input_path}")
            
            # Extract text with enhanced watermark removal
            text = self.watermark_remover.extract_text_from_pdf(input_path)
            logger.info(f"Extracted text length: {len(text)} characters")
            
            # Extract images for integration (disabled for now due to extraction issues)
            images = []  # self.watermark_remover._extract_images_from_pdf(input_path)
            logger.info(f"Extracted {len(images)} images from document")
            
            # Detect SAT structure with image integration
            sat_structure = self._detect_sat_structure_with_images(text, images)
            logger.info(f"Detected SAT structure with images: {sat_structure}")
            
            # Process and format for Word with unified processing
            formatted_text = self._format_sat_for_word_unified(text, sat_structure, images)
            
            return formatted_text
            
        except Exception as e:
            logger.error(f"SAT document processing with unified processing failed: {e}")
            raise
    
    def _detect_sat_structure(self, text: str) -> Dict:
        """Detect SAT document structure and components with enhanced classification"""
        try:
            lines = text.split('\n')
            structure = {
                'sections': [],
                'reading_passages': [],
                'questions': [],
                'multiple_choices': [],
                'document_type': 'unknown',
                'total_questions': 0
            }
            
            current_passage = None
            current_question = None
            question_counter = 0
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # Detect section headers
                if self._matches_pattern(line, self.sat_patterns['section_headers']):
                    structure['sections'].append({
                        'text': line,
                        'line_number': i,
                        'section_type': self._classify_section_type(line)
                    })
                
                # Detect reading passages
                elif self._is_reading_passage_line(line):
                    if current_passage:
                        structure['reading_passages'].append(current_passage)
                    current_passage = {
                        'text': line,
                        'content': [],
                        'line_number': i,
                        'passage_type': 'reading'
                    }
                elif current_passage and not self._is_question_line(line) and not self._is_multiple_choice_line(line):
                    # Add content to current passage if it's not a question or choice
                    if len(line.split()) >= 3:  # Only add substantial content
                        current_passage['content'].append(line)
                
                # Detect questions
                elif self._is_question_line(line):
                    if current_question:
                        # Classify the question before adding
                        try:
                            question_classification = self._classify_question_type(
                                current_question.get('text', '')
                            )
                            current_question['section_type'] = question_classification.get('section_type', 'unknown')
                        except Exception as e:
                            logger.warning(f"Question classification failed: {e}")
                            current_question['section_type'] = 'unknown'
                        
                        # Link passage content to question if available
                        if current_passage and current_passage['content']:
                            current_question['passage_content'] = current_passage['content']
                            current_question['passage_type'] = current_passage.get('passage_type', 'reading')
                        
                        structure['questions'].append(current_question)
                    
                    # Start new question
                    question_counter += 1
                    current_question = {
                        'text': line,
                        'content': [line],
                        'choices': [],
                        'line_number': i,
                        'question_number': question_counter,
                        'images': [],  # Initialize images list
                        'section_type': 'unknown',
                        'passage_content': [],
                        'passage_type': 'none'
                    }
                
                # Detect multiple choice options
                elif self._is_multiple_choice_line(line):
                    if current_question:
                        # Clean up the choice text
                        choice_text = re.sub(r'^[A-D][\)\.\s]*', '', line)
                        choice_text = re.sub(r'\s*[A-Z]+\s*$', '', choice_text)  # Remove trailing letters like "A+"
                        choice_text = choice_text.strip()
                        
                        if choice_text:  # Only add if there's actual content
                            choice_obj = {
                                'text': choice_text,
                                'option': line[0] if line else 'A',
                                'is_correct': False
                            }
                            current_question['choices'].append(choice_obj)
                            
                            # Also add to the global multiple_choices list
                            structure['multiple_choices'].append(choice_obj)
                        continue
                
                # Add content to current passage or question
                else:
                    # Digital SAT: Better content assignment logic
                    if current_passage:
                        # Only add substantial content to passages (not single words/numbers)
                        if len(line.split()) >= 3 or len(line) > 20:
                            current_passage['content'].append(line)
                    elif current_question:
                        # Add content to questions, but be more selective
                        if len(line.split()) >= 2 or len(line) > 10:
                            current_question['content'].append(line)
            
            # Add final passage and question
            if current_passage:
                structure['reading_passages'].append(current_passage)
            if current_question:
                # Classify the final question
                try:
                    question_classification = self._classify_question_type(
                        current_question.get('text', '')
                    )
                    current_question['section_type'] = question_classification.get('section_type', 'unknown')
                except Exception as e:
                    logger.warning(f"Final question classification failed: {e}")
                    current_question['section_type'] = 'unknown'
                
                # Link passage content to final question if available
                if current_passage and current_passage['content']:
                    current_question['passage_content'] = current_passage['content']
                    current_question['passage_type'] = current_passage.get('passage_type', 'reading')
                
                structure['questions'].append(current_question)
            
            # Determine overall document type
            structure['total_questions'] = question_counter
            structure['document_type'] = self._determine_document_type(structure)
            
            # Enhance with AI for better detection
            structure = self._enhance_sat_detection_with_ai(text, structure)
            
            logger.info(f"Detected {len(structure['sections'])} sections, {len(structure['reading_passages'])} passages, {len(structure['questions'])} questions")
            logger.info(f"Document type: {structure['document_type']}")
            return structure
            
        except Exception as e:
            logger.error(f"SAT structure detection failed: {e}")
            return {'sections': [], 'reading_passages': [], 'questions': [], 'document_type': 'unknown', 'total_questions': 0}
    
    def _classify_section_type(self, section_text: str) -> str:
        """Classify section type based on header text"""
        section_lower = section_text.lower()
        if 'math' in section_lower or 'mathematics' in section_lower:
            return 'math'
        elif 'reading' in section_lower or 'evidence' in section_lower:
            return 'reading'
        elif 'writing' in section_lower or 'language' in section_lower:
            return 'writing'
        else:
            return 'general'
    
    def _determine_document_type(self, structure: Dict) -> str:
        """Determine overall document type based on structure analysis"""
        if not structure['questions']:
            return 'unknown'
        
        # Count question types
        math_questions = sum(1 for q in structure['questions'] if q.get('section_type') == 'math')
        reading_questions = sum(1 for q in structure['questions'] if q.get('section_type') == 'reading')
        writing_questions = sum(1 for q in structure['questions'] if q.get('section_type') == 'writing')
        
        total = len(structure['questions'])
        
        if math_questions > total * 0.6:
            return 'math_focused'
        elif reading_questions > total * 0.6:
            return 'reading_focused'
        elif writing_questions > total * 0.6:
            return 'writing_focused'
        elif math_questions > 0 and reading_questions > 0:
            return 'mixed_sat'
        else:
            return 'general'
    
    def _matches_pattern(self, text: str, patterns: List[str]) -> bool:
        """Check if text matches any of the given patterns"""
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _is_question_line(self, line: str) -> bool:
        """Check if a line is a question line (Q1., Q2., 1., 2., etc.)"""
        # Digital SAT question patterns
        question_patterns = [
            r'^Q\d+\.?\s*',  # Q1., Q2., etc.
            r'^\d+\.\s*',    # 1., 2., etc.
            r'^Question\s+\d+[:\s]*',  # Question 1:, Question 2, etc.
        ]
        for pattern in question_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def _is_multiple_choice_line(self, line: str) -> bool:
        """Check if a line is a multiple choice option (A., B., C., D.)"""
        # Digital SAT multiple choice patterns
        choice_patterns = [
            r'^[A-D]\.\s+',  # A., B., C., D.
            r'^[A-D]\)\s+',  # A), B), C), D)
            r'^[A-D]\s+',    # A, B, C, D
        ]
        for pattern in choice_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def _is_reading_passage_line(self, line: str) -> bool:
        """Check if a line is part of a reading passage"""
        # Reading passage indicators
        passage_indicators = [
            r'^[A-Z][^.!?]*[.!?]$',  # Complete sentences
            r'^[A-Z][a-z].*',        # Starts with capital letter
        ]
        
        # Skip if it's a question or choice
        if self._is_question_line(line) or self._is_multiple_choice_line(line):
            return False
        
        # Check if it looks like a passage
        for pattern in passage_indicators:
            if re.match(pattern, line):
                return True
        return False
    
    def _format_sat_for_word(self, text: str, structure: Dict) -> str:
        """Format SAT content for Word document matching the image format"""
        formatted_lines = []
        
        # Process questions with enhanced format matching the image
        question_number = 1
        for question in structure['questions']:
            # Add question number with label (as shown in image)
            formatted_lines.append(f"Question {question_number}")
            formatted_lines.append("")
            
            # Add reading passage content (without label, as shown in image)
            if 'passage_content' in question and question['passage_content']:
                for line in question['passage_content']:
                    if line.strip():  # Only add non-empty lines
                        formatted_lines.append(line.strip())
                formatted_lines.append("")
            
            # Add question text (without label, as shown in image)
            if 'content' in question and question['content']:
                for line in question['content']:
                    if line.strip():  # Only add non-empty lines
                        formatted_lines.append(line.strip())
                formatted_lines.append("")
            
            # Add multiple choice options with A,B,C,D labels (without "multiple choice" label)
            if question['choices']:
                for choice in question['choices']:
                    if choice.get('text', '').strip():
                        # Format as "A. option text" (matching image format)
                        option_label = choice.get('option', '')
                        option_text = choice['text'].strip()
                        if option_label and option_text:
                            formatted_lines.append(f"{option_label}. {option_text}")
                formatted_lines.append("")
            
            # Add clear separation between questions (as shown in image)
            formatted_lines.append("")
            formatted_lines.append("")
            question_number += 1
        
        return '\n'.join(formatted_lines)
    
    def _detect_sat_structure_with_images(self, text: str, images: List[Dict]) -> Dict:
        """Detect SAT structure with image integration and AI enhancement"""
        try:
            # Use AI to enhance structure detection if available
            if self.watermark_remover._has_gemini_api():
                ai_enhanced_structure = self._ai_enhance_sat_structure_detection(text)
                if ai_enhanced_structure:
                    logger.info("AI-enhanced SAT structure detection completed")
                    # Integrate images with AI-enhanced structure
                    for passage in ai_enhanced_structure['reading_passages']:
                        passage['images'] = []
                        for image in images:
                            if self._image_belongs_to_passage(image, passage):
                                passage['images'].append(image)
                    return ai_enhanced_structure
            
            # Fallback to basic structure detection
            structure = self._detect_sat_structure(text)
            
            # Integrate images into reading passages
            for passage in structure['reading_passages']:
                passage['images'] = []
                
                # Find images that belong to this passage
                for image in images:
                    # Simple heuristic: images on the same page as passage content
                    if self._image_belongs_to_passage(image, passage):
                        passage['images'].append(image)
                        logger.info(f"Image {image['index']} integrated into passage: {passage['text']}")
            
            return structure
            
        except Exception as e:
            logger.error(f"AI-enhanced SAT structure detection failed: {e}")
            # Fallback to basic detection
            return self._detect_sat_structure(text)
    
    def _image_belongs_to_passage(self, image: Dict, passage: Dict) -> bool:
        """Determine if an image belongs to a specific passage"""
        # Enhanced heuristic: image is on the same page as passage content
        # or within a reasonable range of the passage
        
        if 'line_number' in passage:
            # Estimate page number from line number (assuming ~50 lines per page)
            estimated_page = passage['line_number'] // 50 + 1
            return image['page'] == estimated_page
        
        return False
    
    def _integrate_images_with_content_enhanced(self, structure: Dict, images: List[Dict]) -> Dict:
        """Enhanced image integration with better matching"""
        # Integrate images into reading passages
        for passage in structure['reading_passages']:
            passage['images'] = []
            
            # Find images that belong to this passage
            for image in images:
                if self._image_belongs_to_passage(image, passage):
                    passage['images'].append(image)
                    logger.info(f"Image {image['index']} integrated into passage: {passage['text']}")
        
        # Integrate images into questions
        for question in structure['questions']:
            question['images'] = []
            
            # Find images that belong to this question
            for image in images:
                if self._image_belongs_to_question(image, question):
                    question['images'].append(image)
                    logger.info(f"Image {image['index']} integrated into question: {question['text']}")
        
        return structure
    
    def _image_belongs_to_question(self, image: Dict, question: Dict) -> bool:
        """Determine if an image belongs to a specific question"""
        if 'line_number' in question:
            # Estimate page number from line number (assuming ~50 lines per page)
            estimated_page = question['line_number'] // 50 + 1
            return image['page'] == estimated_page
        
        return False
    
    def _link_passages_to_questions(self, structure: Dict) -> Dict:
        """Link reading passages to their corresponding questions"""
        for question in structure['questions']:
            question['passage_content'] = []
            
            # Find the most relevant passage for this question
            best_passage = None
            min_distance = float('inf')
            
            for passage in structure['reading_passages']:
                if 'line_number' in passage and 'line_number' in question:
                    distance = abs(passage['line_number'] - question['line_number'])
                    if distance < min_distance and distance < 100:  # Within reasonable range
                        min_distance = distance
                        best_passage = passage
            
            if best_passage:
                question['passage_content'] = best_passage['content']
                logger.info(f"Linked passage to question {question.get('question_number', 'unknown')}")
        
        return structure
    
    def _format_sat_for_word_with_images(self, text: str, structure: Dict, images: List[Dict]) -> str:
        """Format SAT content for Word document with image integration"""
        formatted_lines = []
        
        # Add title
        formatted_lines.append("SAT Practice Test")
        formatted_lines.append("=" * 50)
        formatted_lines.append("")
        
        # Process sections
        for section in structure['sections']:
            formatted_lines.append(f"**SECTION_HEADER:{section['text']}**")
            formatted_lines.append("")
        
        # Process reading passages with images
        for passage in structure['reading_passages']:
            formatted_lines.append(f"**READING_PASSAGE:{passage['text']}**")
            formatted_lines.append("")
            
            # Add passage content
            if 'content' in passage:
                for line in passage['content']:
                    formatted_lines.append(f"PASSAGE_CONTENT:{line}")
                formatted_lines.append("")
            
            # Add images associated with this passage
            if 'images' in passage and passage['images']:
                formatted_lines.append("**PASSAGE_IMAGES:**")
                for image in passage['images']:
                    formatted_lines.append(f"IMAGE_PLACEHOLDER:Page {image['page']}, Image {image['index']} ({image['width']}x{image['height']})")
                formatted_lines.append("")
        
        # Process questions with proper structure
        for question in structure['questions']:
            formatted_lines.append(f"**QUESTION:{question['text']}**")
            
            # Add question content if any
            if 'content' in question:
                for line in question['content']:
                    formatted_lines.append(f"QUESTION_CONTENT:{line}")
            
            # Add multiple choice options with proper formatting
            if question['choices']:
                formatted_lines.append("")
                for choice in question['choices']:
                    formatted_lines.append(f"MULTIPLE_CHOICE:{choice['text']}")
                formatted_lines.append("")
        
        return '\n'.join(formatted_lines)
    
    def _format_sat_for_word_english(self, text: str, structure: Dict, images: List[Dict]) -> str:
        """Format SAT content for Word document with unified English/Math format"""
        formatted_lines = []
        
        # Process questions with unified format for both English and Math
        question_number = 1
        for question in structure['questions']:
            # Add question number with label
            formatted_lines.append(f"Question {question_number}")
            formatted_lines.append("")
            
            # Determine question type for better processing
            question_type = question.get('question_type', 'unknown')
            section_type = question.get('section_type', 'unknown')
            
            # Add reading passage content (for English sections) - without label
            if section_type in ['english', 'reading', 'writing'] and 'passage_content' in question and question['passage_content']:
                for line in question['passage_content']:
                    if line.strip():  # Only add non-empty lines
                        formatted_lines.append(line.strip())
                formatted_lines.append("")
            
            # Add question text (without label) - for both English and Math
            if 'content' in question and question['content']:
                for line in question['content']:
                    if line.strip():  # Only add non-empty lines
                        # For math questions, convert to LaTeX format
                        if section_type in ['math', 'mathematics']:
                            latex_line = self._convert_math_to_latex(line.strip())
                            formatted_lines.append(latex_line)
                        else:
                            formatted_lines.append(line.strip())
                formatted_lines.append("")
            
            # Add multiple choice options or written answer space
            if question['choices']:
                # Multiple choice format (A., B., C., D.)
                for choice in question['choices']:
                    if choice.get('text', '').strip():
                        option_label = choice.get('option', '')
                        option_text = choice['text'].strip()
                        if option_label and option_text:
                            # For math choices, convert to LaTeX format
                            if section_type in ['math', 'mathematics']:
                                latex_choice = self._convert_math_to_latex(option_text)
                                formatted_lines.append(f"{option_label}. {latex_choice}")
                            else:
                                formatted_lines.append(f"{option_label}. {option_text}")
                formatted_lines.append("")
            # Only show multiple choice if detected, otherwise skip answer section entirely
            
            # Add clear separation between questions
            formatted_lines.append("")
            formatted_lines.append("")
            question_number += 1
        
        return '\n'.join(formatted_lines)
    
    def _format_sat_for_word_unified(self, text: str, structure: Dict, images: List[Dict]) -> str:
        """Format SAT content for Word document following the correct RSS format"""
        formatted_lines = []
        
        # Determine if this is English or Math section based on content
        section_type = self._detect_section_type(text, structure)
        
        # Always use MODULE 1 for both English and Math sections
        formatted_lines.append("MODULE 1")
        
        formatted_lines.append("")
        
        # Process questions with correct format
        question_number = 1
        for question in structure['questions']:
            # Add question number (simple format: "1.", "2.", etc.)
            formatted_lines.append(f"{question_number}.")
            formatted_lines.append("")
            
            # Determine question type for better processing
            question_type = question.get('question_type', 'unknown')
            question_section_type = question.get('section_type', section_type)
            
            # Add images if present (no label)
            if 'images' in question and question['images']:
                for image in question['images']:
                    formatted_lines.append(f"[Image: {image.get('description', 'Figure')}]")
                formatted_lines.append("")
            
            # Add reading passage content (for English sections only) - without label
            if question_section_type in ['english', 'reading', 'writing'] and 'passage_content' in question and question['passage_content']:
                for line in question['passage_content']:
                    if line.strip():  # Only add non-empty lines
                        formatted_lines.append(line.strip())
                formatted_lines.append("")
            
            # Add question text (without label) - for both English and Math
            if 'content' in question and question['content']:
                for line in question['content']:
                    if line.strip():  # Only add non-empty lines
                        # For math questions, convert to LaTeX format
                        if question_section_type in ['math', 'mathematics']:
                            latex_line = self._convert_math_to_latex(line.strip())
                            formatted_lines.append(latex_line)
                        else:
                            formatted_lines.append(line.strip())
                formatted_lines.append("")
            
            # Add multiple choice options (A., B., C., D.) - without label
            if question['choices']:
                for choice in question['choices']:
                    if choice.get('text', '').strip():
                        option_label = choice.get('option', '')
                        option_text = choice['text'].strip()
                        if option_label and option_text:
                            # For math choices, convert to LaTeX format
                            if question_section_type in ['math', 'mathematics']:
                                latex_choice = self._convert_math_to_latex(option_text)
                                formatted_lines.append(f"{option_label}. {latex_choice}")
                            else:
                                formatted_lines.append(f"{option_label}. {option_text}")
                formatted_lines.append("")
            
            # Add answer label (auto-label each question)
            # For now, we'll add a placeholder - this should be enhanced with AI detection
            formatted_lines.append("Answer: ")
            formatted_lines.append("")
            
            # Add clear separation between questions
            formatted_lines.append("")
            question_number += 1
        
        return '\n'.join(formatted_lines)
    
    def _detect_section_type(self, text: str, structure: Dict) -> str:
        """Detect if this is an English or Math section based on content analysis"""
        # Count math indicators vs English indicators
        math_indicators = 0
        english_indicators = 0
        
        # Check for math-specific content
        math_patterns = [
            r'\b(equation|formula|graph|function|algebra|geometry|calculus|trigonometry)\b',
            r'[+\-*/=<>≤≥]',  # Math operators
            r'\b\d+\s*[+\-*/]\s*\d+',  # Basic arithmetic
            r'\b(solve|calculate|find|determine)\b.*\b(value|answer|result)\b',
            r'\b(quadratic|linear|polynomial|exponential)\b'
        ]
        
        # Check for English-specific content
        english_patterns = [
            r'\b(reading|passage|text|author|narrator|character)\b',
            r'\b(which choice|best describes|most likely|suggests)\b',
            r'\b(grammar|punctuation|syntax|convention)\b',
            r'\b(complete|logical|precise|word|phrase)\b'
        ]
        
        # Analyze text content
        for pattern in math_patterns:
            math_indicators += len(re.findall(pattern, text, re.IGNORECASE))
        
        for pattern in english_patterns:
            english_indicators += len(re.findall(pattern, text, re.IGNORECASE))
        
        # Analyze question structure
        for question in structure.get('questions', []):
            question_text = ' '.join(question.get('content', []))
            if any(re.search(pattern, question_text, re.IGNORECASE) for pattern in math_patterns):
                math_indicators += 2
            if any(re.search(pattern, question_text, re.IGNORECASE) for pattern in english_patterns):
                english_indicators += 2
        
        # Return section type based on indicators
        if math_indicators > english_indicators:
            return 'math'
        else:
            return 'english'
    
    def _format_sat_for_word_math(self, text: str, structure: Dict, images: List[Dict]) -> str:
        """Format SAT content for Word document with Math section focus (LaTeX conversion)"""
        formatted_lines = []
        
        # Process questions with enhanced format matching the image
        question_number = 1
        for question in structure['questions']:
            # Add question number with label (as shown in image)
            formatted_lines.append(f"Question {question_number}")
            formatted_lines.append("")
            
            # Add reading passage content (if any) - Math section typically has no reading passages
            if 'passage_content' in question and question['passage_content']:
                for line in question['passage_content']:
                    if line.strip():  # Only add non-empty lines
                        # Convert math expressions to LaTeX format
                        latex_line = self._convert_math_to_latex(line.strip())
                        formatted_lines.append(latex_line)
                formatted_lines.append("")
            
            # Add question text (without label, as shown in image)
            if 'content' in question and question['content']:
                for line in question['content']:
                    if line.strip():  # Only add non-empty lines
                        # Convert math expressions to LaTeX format
                        latex_line = self._convert_math_to_latex(line.strip())
                        formatted_lines.append(latex_line)
                formatted_lines.append("")
            
            # Add multiple choice options with A,B,C,D labels (without "multiple choice" label)
            if question['choices']:
                for choice in question['choices']:
                    if choice.get('text', '').strip():
                        # Format as "A. option text" (matching image format)
                        option_label = choice.get('option', '')
                        option_text = choice['text'].strip()
                        if option_label and option_text:
                            # Convert math expressions to LaTeX format
                            latex_choice = self._convert_math_to_latex(option_text)
                            formatted_lines.append(f"{option_label}. {latex_choice}")
                formatted_lines.append("")
            
            # Add clear separation between questions (as shown in image)
            formatted_lines.append("")
            formatted_lines.append("")
            question_number += 1
        
        return '\n'.join(formatted_lines)
    
    def _enhance_sat_detection_with_ai(self, text: str, structure: Dict) -> Dict:
        """Enhance SAT structure detection using AI with improved training for PDF learning"""
        try:
            if 'gemini' not in self.ai_api_keys:
                logger.info("No Gemini API key available for AI enhancement")
                return structure

            import google.generativeai as genai
            genai.configure(api_key=self.ai_api_keys['gemini'])
            model = genai.GenerativeModel('gemini-2.5-flash')

            # Create enhanced prompt for AI training based on SAT format research and image format
            prompt = f"""
            You are an expert SAT test analyzer. Analyze this SAT test content and enhance the structure detection to match the exact format shown in the reference image.

            Text sample: {text[:3000]}...

            Current structure: {structure}

            TARGET FORMAT (matching the image):
            1. Question Number: "Question 1", "Question 2", etc. (with label)
            2. Reading Passage: Long descriptive text (without any label)
            3. Question Text: The actual question (without any label)
            4. Multiple Choice: "A. option text", "B. option text", etc. (without "multiple choice" label)

            SAT Format Requirements:
            - Reading Section: Contains reading passages followed by questions
            - Math Section: Contains math problems with functions, equations, and multiple choice
            - Writing Section: Contains grammar and writing questions

            Please enhance the structure with:
            1. Better question separation and numbering
            2. Accurate reading passage detection (long descriptive text before questions)
            3. Clear question text identification
            4. Proper multiple choice formatting (A, B, C, D with periods)
            5. Math function detection and LaTeX conversion for math sections
            6. Image detection and integration for reading passages

            Focus on learning from this PDF to:
            - Separate each question clearly with proper spacing
            - Identify reading passages as continuous text blocks
            - Extract question text without extra labels
            - Format multiple choice options as "A. text", "B. text", etc.
            - Detect and preserve images in reading passages

            Return enhanced structure as JSON format with improved question separation and format compliance.
            """

            response = model.generate_content(prompt)

            if response.text:
                # Parse AI response and merge with existing structure
                try:
                    import json
                    ai_enhanced = json.loads(response.text)
                    # Merge AI enhancements with existing structure
                    structure.update(ai_enhanced)
                    logger.info("AI enhancement applied successfully with PDF learning for SAT format detection")
                except json.JSONDecodeError:
                    logger.warning("AI response not in valid JSON format")

            return structure

        except Exception as e:
            logger.warning(f"AI enhancement failed: {e}")
            return structure
    
    def _convert_math_to_latex(self, text: str) -> str:
        """Convert math expressions to LaTeX format using pdftolatex approach"""
        try:
            if not text or not isinstance(text, str):
                return text
            
            # Common math patterns to convert to LaTeX
            latex_text = text
            
            # Convert fractions: 1/2 -> \frac{1}{2}
            latex_text = re.sub(r'(\d+)/(\d+)', r'\\frac{\1}{\2}', latex_text)
            
            # Convert exponents: x^2 -> x^{2}, x^y -> x^{y}
            latex_text = re.sub(r'(\w+)\^(\d+)', r'\1^{\2}', latex_text)
            latex_text = re.sub(r'(\w+)\^(\w+)', r'\1^{\2}', latex_text)
            
            # Convert square roots: sqrt(x) -> \sqrt{x}
            latex_text = re.sub(r'sqrt\(([^)]+)\)', r'\\sqrt{\1}', latex_text)
            
            # Convert Greek letters and common symbols
            latex_text = re.sub(r'\bpi\b', r'\\pi', latex_text)
            latex_text = re.sub(r'\btheta\b', r'\\theta', latex_text)
            latex_text = re.sub(r'\balpha\b', r'\\alpha', latex_text)
            latex_text = re.sub(r'\bbeta\b', r'\\beta', latex_text)
            latex_text = re.sub(r'\bgamma\b', r'\\gamma', latex_text)
            latex_text = re.sub(r'\bdelta\b', r'\\delta', latex_text)
            latex_text = re.sub(r'\bepsilon\b', r'\\epsilon', latex_text)
            latex_text = re.sub(r'\blambda\b', r'\\lambda', latex_text)
            latex_text = re.sub(r'\bmu\b', r'\\mu', latex_text)
            latex_text = re.sub(r'\bsigma\b', r'\\sigma', latex_text)
            latex_text = re.sub(r'\btau\b', r'\\tau', latex_text)
            latex_text = re.sub(r'\bphi\b', r'\\phi', latex_text)
            latex_text = re.sub(r'\bomega\b', r'\\omega', latex_text)
            
            # Convert infinity symbol
            latex_text = re.sub(r'\binfinity\b', r'\\infty', latex_text)
            latex_text = re.sub(r'∞', r'\\infty', latex_text)
            
            # Convert mathematical operators
            latex_text = re.sub(r'\btimes\b', r'\\times', latex_text)
            latex_text = re.sub(r'\bdiv\b', r'\\div', latex_text)
            latex_text = re.sub(r'\bplus\b', r'+', latex_text)
            latex_text = re.sub(r'\bminus\b', r'-', latex_text)
            
            # Convert inequalities
            latex_text = re.sub(r'<=', r'\\leq', latex_text)
            latex_text = re.sub(r'>=', r'\\geq', latex_text)
            latex_text = re.sub(r'<', r'<', latex_text)
            latex_text = re.sub(r'>', r'>', latex_text)
            
            # Convert set notation
            latex_text = re.sub(r'\bin\b', r'\\in', latex_text)
            latex_text = re.sub(r'\bnotin\b', r'\\notin', latex_text)
            latex_text = re.sub(r'\bsubset\b', r'\\subset', latex_text)
            latex_text = re.sub(r'\bsupset\b', r'\\supset', latex_text)
            latex_text = re.sub(r'\bunion\b', r'\\cup', latex_text)
            latex_text = re.sub(r'\bintersection\b', r'\\cap', latex_text)
            
            # Convert logical operators
            latex_text = re.sub(r'\band\b', r'\\land', latex_text)
            latex_text = re.sub(r'\bor\b', r'\\lor', latex_text)
            latex_text = re.sub(r'\bnot\b', r'\\neg', latex_text)
            
            # Convert function notation: f(x) -> f(x) (already correct)
            # Convert limits: lim x->0 -> \lim_{x \to 0}
            latex_text = re.sub(r'lim\s+(\w+)\s*->\s*(\w+)', r'\\lim_{\1 \\to \2}', latex_text)
            
            # Convert integrals: int -> \int
            latex_text = re.sub(r'\bint\b', r'\\int', latex_text)
            
            # Convert summations: sum -> \sum
            latex_text = re.sub(r'\bsum\b', r'\\sum', latex_text)
            
            # Convert products: prod -> \prod
            latex_text = re.sub(r'\bprod\b', r'\\prod', latex_text)
            
            # Convert absolute value: |x| -> |x| (already correct)
            
            # Convert matrices: [a b; c d] -> \begin{matrix} a & b \\ c & d \end{matrix}
            matrix_pattern = r'\[([^]]+)\]'
            def replace_matrix(match):
                content = match.group(1)
                if ';' in content:  # Matrix with semicolon separator
                    rows = content.split(';')
                    matrix_content = ' \\\\ '.join([row.strip().replace(' ', ' & ') for row in rows])
                    return f'\\begin{{matrix}} {matrix_content} \\end{{matrix}}'
                return match.group(0)  # Return as-is if not a matrix
            
            latex_text = re.sub(matrix_pattern, replace_matrix, latex_text)
            
            # Convert vectors: <a,b> -> \langle a, b \rangle
            latex_text = re.sub(r'<([^>]+)>', r'\\langle \1 \\rangle', latex_text)
            
            # Convert degrees: 90° -> 90°
            latex_text = re.sub(r'(\d+)°', r'\1°', latex_text)
            
            # Convert percentages: 50% -> 50\%
            latex_text = re.sub(r'(\d+)%', r'\1\\%', latex_text)
            
            # Convert common functions
            latex_text = re.sub(r'\bsin\b', r'\\sin', latex_text)
            latex_text = re.sub(r'\bcos\b', r'\\cos', latex_text)
            latex_text = re.sub(r'\btan\b', r'\\tan', latex_text)
            latex_text = re.sub(r'\blog\b', r'\\log', latex_text)
            latex_text = re.sub(r'\bln\b', r'\\ln', latex_text)
            latex_text = re.sub(r'\bexp\b', r'\\exp', latex_text)
            
            # Convert natural log: ln(x) -> \ln(x)
            latex_text = re.sub(r'ln\(([^)]+)\)', r'\\ln(\1)', latex_text)
            
            # Convert log base: log_2(x) -> \log_2(x)
            latex_text = re.sub(r'log_(\w+)\(([^)]+)\)', r'\\log_{\1}(\2)', latex_text)
            
            # Convert complex expressions with parentheses
            # Handle nested expressions like f(x) = 2x + 1
            latex_text = re.sub(r'f\((\w+)\)\s*=\s*([^=]+)', r'f(\1) = \2', latex_text)
            
            # Convert domain and range notation: [a,b] -> [a,b] (already correct)
            # Convert open intervals: (a,b) -> (a,b) (already correct)
            
            # Convert complex numbers: a + bi -> a + bi (already correct)
            latex_text = re.sub(r'(\d+)\s*\+\s*(\d+)i', r'\1 + \2i', latex_text)
            
            # Convert scientific notation: 1.5e10 -> 1.5 \times 10^{10}
            latex_text = re.sub(r'(\d+\.?\d*)[eE](\d+)', r'\1 \\times 10^{\2}', latex_text)
            
            # Clean up multiple spaces
            latex_text = re.sub(r'\s+', ' ', latex_text).strip()
            
            logger.info(f"Converted math expression: '{text}' -> '{latex_text}'")
            return latex_text
            
        except Exception as e:
            logger.warning(f"LaTeX conversion failed for '{text}': {e}")
            return text
    

    
    def _classify_section_type_from_content(self, content: List[str]) -> str:
        """Classify section type based on content"""
        content_text = ' '.join(content).lower()
        
        # Math indicators
        if any(word in content_text for word in ['equation', 'function', 'graph', 'algebra', 'geometry', 'trigonometry', 'solve', 'calculate']):
            return 'math'
        
        # English indicators
        if any(word in content_text for word in ['passage', 'author', 'main idea', 'purpose', 'tone', 'grammar', 'punctuation']):
            return 'english'
        
        return 'unknown'


class DocumentProcessor:
    """Enhanced document processing with SAT support"""
    

    def __init__(self):
        self.watermark_remover = EnhancedWatermarkRemover()
        self.sat_processor = SATDocumentProcessor()
        self.advanced_processor = AdvancedSATProcessor()  # New advanced processor
        logger.info("Enhanced document processor initialized with SAT support")
    
    def process_document_advanced(self, input_path: str) -> Tuple[str, Dict]:
        """Process document using the advanced SAT processor"""
        try:
            logger.info(f"Processing document with advanced processor: {input_path}")
            
            # Use the advanced processor
            formatted_text, structure = self.advanced_processor.process_pdf(input_path)
            
            # Create output filename
            original_filename = Path(input_path).name
            base_name = Path(input_path).stem
            output_filename = f"{base_name}_Advanced.docx"
            output_path = Path("processed") / output_filename
            
            # Ensure processed directory exists
            Path("processed").mkdir(exist_ok=True)
            
            # Create Word document
            self._create_advanced_word_document(formatted_text, structure, output_path)
            
            # Prepare stats
            stats = {
                'questions': len(structure['questions']),
                'passages': len(structure['passages']),
                'math_functions': len(structure['math_functions']),
                'total_questions': structure['total_questions'],
                'file_size': output_path.stat().st_size if output_path.exists() else 0
            }
            
            logger.info(f"Advanced processing completed: {stats['questions']} questions, {stats['passages']} passages")
            return str(output_path), stats
            
        except Exception as e:
            logger.error(f"Advanced processing failed: {e}")
            raise
    
    def _create_advanced_word_document(self, formatted_text: str, structure: Dict, output_path: Path):
        """Create a Word document with proper formatting using the advanced processor output"""
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Process the formatted text
        lines = formatted_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if it's a question header
            if line.startswith('Question '):
                # Add question header
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
                
            # Check if it's a multiple choice option
            elif re.match(r'^[A-D]\.\s+', line):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Indent the option
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(line)
                
            # Regular content
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(line)
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Advanced Word document saved: {output_path}")
    
    def process_document(self, file_path: str, section_type: str = 'english') -> Tuple[str, Dict]:
        """Process PDF document with unified English/Math processing and watermark removal"""
        try:
            logger.info(f"Processing {file_path} with unified section processing")
            
            # Check if this is a SAT document first (for both English and Math sections)
            is_sat = self._is_sat_document(file_path)
            if is_sat:
                logger.info("SAT document detected - using unified specialized processor")
                word_path = self._convert_sat_to_word_unified(file_path, file_path)
                section_type = 'sat_unified'  # Mark as unified SAT processing
            else:
                # Direct conversion from original PDF with watermark removal in text processing
                word_path = self._convert_to_word_unified(file_path, file_path)
                section_type = 'unified'
            
            metadata = {
                'success': True,
                'input_file': Path(file_path).name,
                'output_file': Path(word_path).name,
                'section_type': section_type,
                'processing_time': time.time(),
                'watermark_removal': 'completed',
                'word_conversion': 'completed',
                'sat_detection': 'completed' if is_sat else 'not_applicable'
            }
            
            return word_path, metadata
            
        except Exception as e:
            logger.error(f"Processing failed: {str(e)}")
            logger.error(traceback.format_exc())
            raise
    
    def _is_sat_document(self, file_path: str) -> bool:
        """Detect if document is SAT format"""
        try:
            # Quick check of first few pages for SAT indicators
            text = self.watermark_remover.extract_text_from_pdf(file_path, max_pages=3)
            
            # If text extraction failed (likely image-based PDF), use filename as hint
            if len(text) < 10:  # Very little text extracted
                filename = Path(file_path).name.lower()
                sat_filename_indicators = ['sat', 'test', 'exam', 'practice']
                if any(indicator in filename for indicator in sat_filename_indicators):
                    logger.info(f"SAT detection based on filename: {filename}")
                    return True
            
            return self._is_sat_document_from_text(text)
            
        except Exception as e:
            logger.warning(f"SAT detection failed: {e}")
            return False
    
    def _is_sat_document_from_text(self, text: str) -> bool:
        """Detect if text content is SAT format"""
        sat_indicators = [
            'sat', 'scholastic aptitude test', 'college board',
            'reading passage', 'evidence-based reading',
            'questions are based on the following passage',
            'multiple choice', 'a)', 'b)', 'c)', 'd)',
            'passage', 'question', 'section'
        ]
        
        # More flexible SAT detection patterns
        sat_patterns = [
            r'[A-D]\.\s*[A-Za-z]',  # A. choice, B. choice, etc.
            r'[A-D]\s*[A-Za-z]',     # A choice, B choice, etc.
            r'\d+\s*[A-Za-z]',       # 52 Which choice..., 53 Which choice...
            r'which choice completes', # "Which choice completes the text..."
            r'what is the value',      # Math questions
            r'the graph of the function', # Math questions
            r'section.*module.*math',   # Section headers
            r'reading.*writing.*math'   # Section headers
        ]
        
        text_lower = text.lower()
        
        # Count basic indicators
        indicator_count = sum(1 for indicator in sat_indicators if indicator in text_lower)
        
        # Count pattern matches
        import re
        pattern_count = sum(1 for pattern in sat_patterns if re.search(pattern, text, re.IGNORECASE))
        
        # More flexible detection - if we find multiple choice patterns or question patterns
        total_score = indicator_count + pattern_count
        
        print(f"SAT Detection Debug: indicators={indicator_count}, patterns={pattern_count}, total={total_score}")
        
        # Lower threshold for detection
        return total_score >= 2
    
    def _convert_sat_to_word_unified(self, pdf_path: str, original_filename: str) -> str:
        """Convert SAT PDF to Word document with unified English/Math processing"""
        logger.info("Converting SAT document to Word with unified processing...")
        
        try:
            # Process SAT document through unified processor
            formatted_text = self.sat_processor.process_sat_document_unified(pdf_path)
            
            # Create Word document
            word_doc = Document()
            
            # Set document properties
            word_doc.core_properties.title = f"SAT Practice Test - {Path(original_filename).stem}"
            word_doc.core_properties.author = "SAT Document Processor"
            
            # NO HEADERS - Start directly with questions
            
            # Process formatted text with unified SAT format (Question {number} format)
            lines = formatted_text.split('\n')
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Process new SAT format: Question {number}
                if line.startswith("Question "):
                    current_section = "question"
                    # Add question number as heading
                    heading = word_doc.add_heading(line, level=2)
                    heading.style.font.bold = True
                    heading.style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                    heading.paragraph_format.space_after = Pt(12)
                    continue
                
                # Process multiple choice options (A., B., C., D.)
                if line.startswith(('A.', 'B.', 'C.', 'D.')):
                    current_section = "options"
                    para = word_doc.add_paragraph(line)
                    para.paragraph_format.space_after = Pt(6)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.style.font.size = Pt(11)
                    continue
                
                # Process regular content (reading passages, question text)
                if line and not line.startswith(('Question ', 'A.', 'B.', 'C.', 'D.')):
                    # This is reading passage content or question text
                    para = word_doc.add_paragraph(line)
                    para.paragraph_format.space_after = Pt(6)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.style.font.size = Pt(11)
                    continue
            
            # Save Word document
            output_path = os.path.join(PROCESSED_FOLDER, f"{Path(original_filename).stem}_SAT_Unified.docx")
            word_doc.save(output_path)
            
            logger.info(f"Unified SAT Word document saved: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Unified SAT to Word conversion failed: {e}")
            raise
    
    def _convert_to_word_unified(self, pdf_path: str, original_filename: str) -> str:
        """Convert PDF to Word document with unified processing (English/Math)"""
        logger.info("Converting to Word document with unified processing...")
        
        try:
            # Open the cleaned PDF
            doc = fitz.open(pdf_path)
            
            # Create Word document
            word_doc = Document()
            
            # Set document properties
            word_doc.core_properties.title = f"Converted Document - {Path(original_filename).stem}"
            word_doc.core_properties.author = "PDF Watermark Remover - Unified Processing"
            
            # Add title
            title = word_doc.add_heading("Converted Document - Unified Processing", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add subtitle
            subtitle = word_doc.add_paragraph(f"Source: {Path(original_filename).name}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add watermark removal info
            info = word_doc.add_paragraph("Watermarks removed using advanced detection algorithms")
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Add page separator
                if page_num > 0:
                    word_doc.add_page_break()
                
                # Add page header
                page_header = word_doc.add_heading(f"Page {page_num + 1}", level=1)
                
                # Extract text from page
                text = page.get_text()
                
                # Remove watermarks from text
                text = self.watermark_remover._remove_watermarks_enhanced(text)
                
                # Clean up text for English section
                text = self._clean_text_english(text)
                
                # Split into paragraphs and add to Word
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Check if this looks like a heading
                        if len(para_text.strip()) < 100 and para_text.strip().isupper():
                            word_doc.add_heading(para_text.strip(), level=2)
                        else:
                            word_doc.add_paragraph(para_text.strip())
                
                # Add page metadata
                word_count = len(text.split())
                meta_para = word_doc.add_paragraph(f"Words on this page: {word_count}")
                meta_para.style = 'Quote'
            
            # Close PDF
            doc.close()
            
            # Save Word document
            output_filename = f"converted_english_{Path(original_filename).stem}.docx"
            output_path = os.path.join(PROCESSED_FOLDER, output_filename)
            
            # Test the document before saving
            try:
                word_doc.save(output_path)
                
                # Verify the file can be opened
                test_doc = Document(output_path)
                # Document objects don't have close() method, just verify it loaded
                
                logger.info(f"Word document saved and verified: {output_path}")
                return output_path
                
            except Exception as e:
                logger.error(f"Word document save/verification failed: {e}")
                raise Exception(f"Word document generation failed: {e}")
            
        except Exception as e:
            logger.error(f"Word conversion failed: {str(e)}")
            raise
    
    def _convert_to_word_math(self, pdf_path: str, original_filename: str) -> str:
        """Convert PDF to Word document for Math section (LaTeX conversion)"""
        logger.info("Converting to Word document (Math section)...")
        
        try:
            # Open the cleaned PDF
            doc = fitz.open(pdf_path)
            
            # Create Word document
            word_doc = Document()
            
            # Set document properties
            word_doc.core_properties.title = f"Converted Document - {Path(original_filename).stem}"
            word_doc.core_properties.author = "PDF Watermark Remover - Math Section"
            
            # Add title
            title = word_doc.add_heading("Converted Document - Math Section", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add subtitle
            subtitle = word_doc.add_paragraph(f"Source: {Path(original_filename).stem}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add math section info
            info = word_doc.add_paragraph("Math expressions converted to LaTeX format")
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Add page separator
                if page_num > 0:
                    word_doc.add_page_break()
                
                # Add page header
                page_header = word_doc.add_heading(f"Page {page_num + 1}", level=1)
                
                # Extract text from page
                text = page.get_text()
                
                # Remove watermarks from text
                text = self.watermark_remover._remove_watermarks_enhanced(text)
                
                # Convert math expressions to LaTeX
                text = self._convert_math_to_latex(text)
                
                # Clean up text for Math section
                text = self._clean_text_math(text)
                
                # Split into paragraphs and add to Word
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Check if this looks like a heading
                        if len(para_text.strip()) < 100 and para_text.strip().isupper():
                            word_doc.add_heading(para_text.strip(), level=2)
                        else:
                            word_doc.add_paragraph(para_text.strip())
                
                # Add page metadata
                word_count = len(text.split())
                meta_para = word_doc.add_paragraph(f"Words on this page: {word_count}")
                meta_para.style = 'Quote'
            
            # Close PDF
            doc.close()
            
            # Save Word document
            output_filename = f"converted_math_{Path(original_filename).stem}.docx"
            output_path = os.path.join(PROCESSED_FOLDER, output_filename)
            
            # Test the document before saving
            try:
                word_doc.save(output_path)
                
                # Verify the file can be opened
                test_doc = Document(output_path)
                # Document objects don't have close() method, just verify it loaded
                
                logger.info(f"Word document saved and verified: {output_path}")
                return output_path
                
            except Exception as e:
                logger.error(f"Word document save/verification failed: {e}")
                raise Exception(f"Word document generation failed: {e}")
            
        except Exception as e:
            logger.error(f"Word conversion failed: {str(e)}")
            raise
    
    def _convert_math_to_latex(self, text: str) -> str:
        """Enhanced math to LaTeX conversion inspired by LaTeX-OCR repository"""
        # Advanced math conversion patterns
        conversions = [
            # Greek letters (process first to avoid conflicts)
            (r'\balpha\b', r'\\alpha'),
            (r'\bbeta\b', r'\\beta'),
            (r'\bgamma\b', r'\\gamma'),
            (r'\bdelta\b', r'\\delta'),
            (r'\bepsilon\b', r'\\epsilon'),
            (r'\bzeta\b', r'\\zeta'),
            (r'\beta\b', r'\\beta'),
            (r'\btheta\b', r'\\theta'),
            (r'\biota\b', r'\\iota'),
            (r'\bkappa\b', r'\\kappa'),
            (r'\blambda\b', r'\\lambda'),
            (r'\bmu\b', r'\\mu'),
            (r'\bnu\b', r'\\nu'),
            (r'\bxi\b', r'\\xi'),
            (r'\bpi\b', r'\\pi'),
            (r'\brho\b', r'\\rho'),
            (r'\bsigma\b', r'\\sigma'),
            (r'\btau\b', r'\\tau'),
            (r'\bupsilon\b', r'\\upsilon'),
            (r'\bphi\b', r'\\phi'),
            (r'\bchi\b', r'\\chi'),
            (r'\bpsi\b', r'\\psi'),
            (r'\bomega\b', r'\\omega'),
            
            # Fractions and ratios
            (r'(\w+)/(\w+)', r'\\frac{\1}{\2}'),
            (r'(\d+)/(\d+)', r'\\frac{\1}{\2}'),
            
            # Exponents and powers
            (r'(\w+)\^(\d+)', r'\1^{\2}'),
            (r'(\w+)\^(\w+)', r'\1^{\2}'),
            (r'(\d+)\^(\d+)', r'\1^{\2}'),
            
            # Roots and radicals
            (r'sqrt\(([^)]+)\)', r'\\sqrt{\1}'),
            (r'√\(([^)]+)\)', r'\\sqrt{\1}'),
            (r'cbrt\(([^)]+)\)', r'\\sqrt[3]{\1}'),
            (r'root\(([^)]+)\)', r'\\sqrt{\1}'),
            
            # Mathematical operators
            (r'integral', r'\\int'),
            (r'sum', r'\\sum'),
            (r'product', r'\\prod'),
            (r'infinity', r'\\infty'),
            (r'partial', r'\\partial'),
            (r'nabla', r'\\nabla'),
            (r'forall', r'\\forall'),
            (r'exists', r'\\exists'),
            (r'nexists', r'\\nexists'),
            (r'in', r'\\in'),
            (r'notin', r'\\notin'),
            (r'subset', r'\\subset'),
            (r'supset', r'\\supset'),
            (r'subseteq', r'\\subseteq'),
            (r'supseteq', r'\\supseteq'),
            
            # Comparison operators
            (r'<=', r'\\leq'),
            (r'>=', r'\\geq'),
            (r'!=', r'\\neq'),
            (r'approx', r'\\approx'),
            (r'equiv', r'\\equiv'),
            (r'propto', r'\\propto'),
            (r'sim', r'\\sim'),
            (r'cong', r'\\cong'),
            
            # Arrows
            (r'->', r'\\rightarrow'),
            (r'<-', r'\\leftarrow'),
            (r'<->', r'\\leftrightarrow'),
            (r'=>', r'\\Rightarrow'),
            (r'<=', r'\\Leftarrow'),
            (r'<=>', r'\\Leftrightarrow'),
            (r'mapsto', r'\\mapsto'),
            
            # Subscripts and superscripts
            (r'(\w+)_(\w+)', r'\1_{\2}'),
            (r'(\w+)_(\d+)', r'\1_{\2}'),
            (r'(\d+)_(\w+)', r'\1_{\2}'),
            (r'(\d+)_(\d+)', r'\1_{\2}'),
            
            # Common mathematical functions
            (r'sin\(', r'\\sin('),
            (r'cos\(', r'\\cos('),
            (r'tan\(', r'\\tan('),
            (r'log\(', r'\\log('),
            (r'ln\(', r'\\ln('),
            (r'exp\(', r'\\exp('),
            (r'lim\(', r'\\lim('),
            (r'max\(', r'\\max('),
            (r'min\(', r'\\min('),
            
            # Sets and logic
            (r'emptyset', r'\\emptyset'),
            (r'mathbb\{R\}', r'\\mathbb{R}'),
            (r'mathbb\{N\}', r'\\mathbb{N}'),
            (r'mathbb\{Z\}', r'\\mathbb{Z}'),
            (r'mathbb\{Q\}', r'\\mathbb{Q}'),
            (r'mathbb\{C\}', r'\\mathbb{C}'),
        ]
        
        processed_text = text
        
        # Apply conversions
        for pattern, replacement in conversions:
            processed_text = re.sub(pattern, replacement, processed_text, flags=re.IGNORECASE)
        
        # Enhanced LaTeX expression detection and wrapping
        # Look for LaTeX commands and wrap them properly
        latex_patterns = [
            (r'\\[a-zA-Z]+(\{[^}]*\})?', r'$$\g<0>$$'),  # Basic LaTeX commands
            (r'\\frac\{[^}]*\}\{[^}]*\}', r'$$\g<0>$$'),  # Fractions
            (r'\\sqrt\{[^}]*\}', r'$$\g<0>$$'),  # Square roots
            (r'\\int[^$]*', r'$$\g<0>$$'),  # Integrals
            (r'\\sum[^$]*', r'$$\g<0>$$'),  # Sums
        ]
        
        for pattern, replacement in latex_patterns:
            processed_text = re.sub(pattern, replacement, processed_text)
        
        # Clean up multiple dollar signs
        processed_text = re.sub(r'\$\$\$\$', '$$', processed_text)
        
        return processed_text
    
    def _clean_text_english(self, text: str) -> str:
        """Clean and format text for English section Word document"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR issues
        replacements = [
            ('||', 'll'), ('|/', 'll'), ('0O', '0'), ('O0', '0'),
            ('1l', 'll'), ('l1', 'll'), ('5S', 'S'), ('S5', 'S'),
            ('rn', 'm'), ('cl', 'd'), ('vv', 'w'), ('nn', 'm')
        ]
        
        for old, new in replacements:
            text = text.replace(old, new)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def _clean_text_math(self, text: str) -> str:
        """Clean and format text for Math section Word document"""
        # Remove extra whitespace but preserve math formatting
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR issues for math
        replacements = [
            ('||', 'll'), ('|/', 'll'), ('0O', '0'), ('O0', '0'),
            ('1l', 'll'), ('l1', 'll'), ('5S', 'S'), ('S5', 'S'),
            ('rn', 'm'), ('cl', 'd'), ('vv', 'w'), ('nn', 'm')
        ]
        
        for old, new in replacements:
            text = text.replace(old, new)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()

# Initialize processor
processor = None
app_startup_time = None

def get_processor():
    """Lazy initialization of processor for faster startup"""
    global processor
    if processor is None:
        processor = DocumentProcessor()
    return processor

@app.route('/')
def root():
    """Root endpoint - redirect to home page for web interface"""
    return redirect('/home')

@app.route('/startup')
def startup():
    """Startup endpoint for Railway - responds immediately"""
    return jsonify({
        'status': 'ready',
        'message': 'App startup complete',
        'timestamp': time.time(),
        'startup_time': app_startup_time,
        'uptime': time.time() - app_startup_time if app_startup_time else 0
    })

@app.route('/ready')
def ready():
    """Readiness probe for Railway - checks if app is fully operational"""
    try:
        # Quick test of core functionality
        test_processor = get_processor()
        return jsonify({
            'status': 'ready',
            'message': 'Application fully operational',
            'timestamp': time.time(),
            'startup_time': app_startup_time,
            'uptime': time.time() - app_startup_time if app_startup_time else 0,
            'processor_status': 'initialized',
            'railway_ready': True
        })
    except Exception as e:
        return jsonify({
            'status': 'not_ready',
            'message': f'Application not ready: {str(e)}',
            'timestamp': time.time(),
            'error': str(e)
        }), 503

@app.route('/live')
def live():
    """Liveness probe for Railway - responds immediately"""
    return jsonify({
        'status': 'alive',
        'message': 'Application is running',
        'timestamp': time.time()
    })

@app.route('/api-key', methods=['POST'])
def update_api_key():
    """Update API key for AI enhancement"""
    try:
        data = request.get_json()
        api_key = data.get('api_key', '').strip()
        provider = data.get('provider', 'gemini')
        
        if not api_key:
            return jsonify({'success': False, 'error': 'API key is required'}), 400
        
        # Store globally and in watermark remover
        global GEMINI_API_KEY
        GEMINI_API_KEY = api_key
        
        # Update watermark remover instance
        watermark_remover = get_processor().watermark_remover
        watermark_remover.set_api_key(provider, api_key)
        
        logger.info(f"API key updated for {provider}")
        return jsonify({
            'success': True, 
            'message': f'API key updated successfully for {provider}'
        })
        
    except Exception as e:
        logger.error(f"API key update failed: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    """Submit user feedback for AI training"""
    try:
        data = request.get_json()
        feedback = data.get('feedback', '').strip()
        timestamp = data.get('timestamp', '')
        
        if not feedback:
            return jsonify({'error': 'Feedback is required'}), 400
        
        # Log feedback for AI training
        logger.info(f"User feedback received: {feedback}")
        
        # Store feedback for AI training (in a real implementation, you'd save to database)
        feedback_data = {
            'feedback': feedback,
            'timestamp': timestamp,
            'user_agent': request.headers.get('User-Agent', ''),
            'ip_address': request.remote_addr,
            'training_type': 'sat_format_improvement'
        }
        
        # Save feedback to file for AI training analysis
        try:
            import json
            import os
            
            feedback_file = 'training_feedback.json'
            feedback_list = []
            
            # Load existing feedback if file exists
            if os.path.exists(feedback_file):
                with open(feedback_file, 'r', encoding='utf-8') as f:
                    feedback_list = json.load(f)
            
            # Add new feedback
            feedback_list.append(feedback_data)
            
            # Save updated feedback
            with open(feedback_file, 'w', encoding='utf-8') as f:
                json.dump(feedback_list, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Feedback saved to {feedback_file} for AI training")
            
        except Exception as e:
            logger.warning(f"Failed to save feedback to file: {e}")
        
        logger.info(f"Feedback data: {feedback_data}")
        
        return jsonify({
            'message': 'Feedback submitted successfully',
            'status': 'received'
        })
        
    except Exception as e:
        logger.error(f"Feedback submission failed: {e}")
        return jsonify({'error': 'Failed to submit feedback'}), 500

@app.route('/home')
def home():
    """Simple home page with two sections"""
    html_content = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>PDF Watermark Remover - Two Sections</title>
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }

            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: #f5f7fa;
                min-height: 100vh;
                padding: 20px;
            }

            .container {
                max-width: 800px;
                margin: 0 auto;
                background: white;
                border-radius: 20px;
                box-shadow: 0 20px 40px rgba(0, 0, 0, 0.1);
                padding: 40px;
            }

            h1 {
                color: #333;
                text-align: center;
                margin-bottom: 10px;
                font-size: 2.2em;
            }

            .subtitle {
                color: #666;
                text-align: center;
                margin-bottom: 40px;
                font-size: 1.1em;
            }

            .api-key-section {
                background: #f8f9fa;
                border-radius: 15px;
                padding: 25px;
                margin-bottom: 30px;
                border: 1px solid #e9ecef;
            }

            .api-key-section h3 {
                color: #333;
                margin-bottom: 10px;
                font-size: 1.3em;
            }

            .api-key-section p {
                color: #666;
                margin-bottom: 20px;
                font-size: 0.95em;
            }

            .api-input-group {
                display: flex;
                gap: 15px;
                margin-bottom: 15px;
            }

            .api-key-input {
                flex: 1;
                padding: 12px 15px;
                border: 2px solid #ddd;
                border-radius: 8px;
                font-size: 1em;
                transition: border-color 0.3s ease;
            }

            .api-key-input:focus {
                outline: none;
                border-color: #4CAF50;
            }

            .api-key-btn {
                background: linear-gradient(45deg, #4CAF50, #45a049);
                color: white;
                border: none;
                padding: 12px 20px;
                border-radius: 8px;
                font-size: 1em;
                font-weight: 600;
                cursor: pointer;
                transition: all 0.3s ease;
                white-space: nowrap;
            }

            .api-key-btn:hover {
                transform: translateY(-2px);
                box-shadow: 0 5px 15px rgba(76, 175, 80, 0.3);
            }

            .api-key-status {
                min-height: 20px;
                font-size: 0.9em;
            }

            .api-key-status.success {
                color: #155724;
            }

            .api-key-status.error {
                color: #721c24;
            }

            /* Training Section Styles */
            .training-section {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 20px;
                border-radius: 12px;
                margin: 20px 0;
                box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);
            }

            .training-section h3 {
                margin: 0 0 10px 0;
                font-size: 18px;
            }

            .training-section p {
                margin: 0 0 15px 0;
                opacity: 0.9;
            }

            .training-controls {
                display: flex;
                gap: 10px;
                flex-wrap: wrap;
            }

            .training-btn, .feedback-btn {
                background: rgba(255, 255, 255, 0.2);
                color: white;
                border: 2px solid rgba(255, 255, 255, 0.3);
                padding: 10px 20px;
                border-radius: 8px;
                cursor: pointer;
                font-size: 14px;
                font-weight: 600;
                transition: all 0.3s ease;
            }

            .training-btn:hover, .feedback-btn:hover {
                background: rgba(255, 255, 255, 0.3);
                border-color: rgba(255, 255, 255, 0.5);
                transform: translateY(-2px);
            }

            .training-status {
                margin-top: 15px;
                padding: 10px;
                border-radius: 6px;
                font-size: 14px;
                text-align: center;
                background: rgba(255, 255, 255, 0.1);
            }

            .sections {
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 30px;
                margin-bottom: 40px;
            }



            .section {
                background: #f8f9fa;
                border-radius: 15px;
                padding: 30px;
                text-align: center;
                border: 2px solid transparent;
                transition: all 0.3s ease;
                cursor: pointer;
            }

            .section:hover {
                border-color: #4CAF50;
                transform: translateY(-5px);
                box-shadow: 0 10px 25px rgba(0, 0, 0, 0.1);
            }

            .section.active {
                border-color: #4CAF50;
                background: #f0fff0;
            }

            .section-icon {
                font-size: 3em;
                margin-bottom: 20px;
            }

            .section-title {
                font-size: 1.5em;
                font-weight: 600;
                color: #333;
                margin-bottom: 15px;
            }

                        .section-desc {
                color: #666;
                line-height: 1.5;
            }

            .drop-zone {
                border: 3px dashed #ddd;
                border-radius: 15px;
                padding: 60px 20px;
                margin-bottom: 30px;
                cursor: pointer;
                transition: all 0.3s ease;
                background: #fafafa;
                text-align: center;
            }

            .drop-zone:hover,
            .drop-zone.dragover {
                border-color: #4CAF50;
                background: #f0fff0;
                transform: scale(1.02);
            }

            .drop-icon {
                font-size: 4em;
                color: #ddd;
                margin-bottom: 20px;
            }

            .drop-zone.dragover .drop-icon {
                color: #4CAF50;
            }

            .drop-text {
                font-size: 1.3em;
                color: #666;
                margin-bottom: 10px;
            }

            .drop-subtext {
                color: #999;
                font-size: 0.9em;
            }

            .file-input {
                display: none;
            }

            .progress-section {
                display: none;
                margin: 30px 0;
            }

            .progress-bar {
                background: #f0f0f0;
                border-radius: 10px;
                overflow: hidden;
                height: 20px;
                margin-bottom: 15px;
            }

            .progress-fill {
                background: linear-gradient(90deg, #4CAF50, #45a049);
                height: 100%;
                width: 0%;
                transition: width 0.3s ease;
                border-radius: 10px;
            }

            .status-text {
                color: #666;
                font-size: 1.1em;
                margin-bottom: 20px;
            }

            .download-btn {
                display: none;
                background: linear-gradient(45deg, #4CAF50, #45a049);
                color: white;
                border: none;
                padding: 15px 40px;
                border-radius: 25px;
                font-size: 1.2em;
                font-weight: 600;
                cursor: pointer;
                transition: all 0.3s ease;
                box-shadow: 0 5px 15px rgba(76, 175, 80, 0.3);
                margin: 0 auto;
            }

            .download-btn:hover {
                transform: translateY(-2px);
                box-shadow: 0 10px 25px rgba(76, 175, 80, 0.4);
            }

            .success-message {
                display: none;
                background: #d4edda;
                color: #155724;
                border: 1px solid #c3e6cb;
                border-radius: 10px;
                padding: 20px;
                margin-bottom: 20px;
                text-align: center;
            }

            .error-message {
                display: none;
                background: #f8d7da;
                color: #721c24;
                border: 1px solid #f5c6cb;
                border-radius: 10px;
                padding: 20px;
                margin-bottom: 20px;
                text-align: center;
            }

            @media (max-width: 768px) {
                .sections {
                    grid-template-columns: 1fr;
                    gap: 20px;
                }
                
                .container {
                    padding: 30px 20px;
                }
                
                .training-section {
                    padding: 20px;
                }
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>📄 PDF Watermark Remover</h1>
            <p class="subtitle">Two sections: English (text) and Math (LaTeX)</p>

                    <!-- API Key Input Section -->
        <div class="api-key-section">
            <h3>🔑 AI Enhancement Setup</h3>
            <p>Enter your Gemini API key for enhanced watermark removal and AI-powered processing:</p>
            <div class="api-input-group">
                <input type="password" id="apiKeyInput" placeholder="Enter your Gemini API key" class="api-key-input">
                <button onclick="updateApiKey()" class="api-key-btn">Set API Key</button>
            </div>
            <div id="apiKeyStatus" class="api-key-status"></div>
        </div>

        <!-- AI Training Section -->
        <div class="training-section">
            <h3>🤖 AI Training & Feedback</h3>
            <p>Help improve the AI by providing feedback on watermark detection and format recognition:</p>
            <div class="training-controls">
                <button onclick="startTraining()" class="training-btn">Start AI Training Session</button>
                <button onclick="provideFeedback()" class="feedback-btn">Provide Feedback</button>
            </div>
            <div id="trainingStatus" class="training-status"></div>
        </div>

            <div class="sections">
                <div class="section" id="englishSection" onclick="selectSection('english')">
                    <div class="section-icon">📝</div>
                    <div class="section-title">English Section</div>
                    <div class="section-desc">
                        Remove watermarks and convert to Word with clean text formatting.
                        <strong>Auto-detects SAT documents</strong> for enhanced processing.
                        Perfect for documents, reports, and SAT practice tests.
                    </div>
                </div>

                <div class="section" id="mathSection" onclick="selectSection('math')">
                    <div class="section-icon">🧮</div>
                    <div class="section-title">Math Section</div>
                    <div class="section-desc">
                        Auto-detect math functions and convert to Word with LaTeX content.
                        <strong>Auto-detects SAT documents</strong> for enhanced processing.
                        Ideal for academic papers, equations, and SAT math sections.
                    </div>
                </div>
            </div>



            <div class="drop-zone" id="dropZone">
                <div class="drop-icon">☁️</div>
                <div class="drop-text">Drop PDF file here or click to browse</div>
                <div class="drop-subtext">Select a section above first</div>
                <input type="file" id="fileInput" class="file-input" accept=".pdf">
            </div>

            <div class="progress-section" id="progressSection">
                <div class="progress-bar">
                    <div class="progress-fill" id="progressFill"></div>
                </div>
                <div class="status-text" id="statusText">Processing...</div>
            </div>

            <div class="success-message" id="successMessage">
                <h3>✅ Processing Complete!</h3>
                <p>Your document has been successfully converted. Click download to get your Word file.</p>
            </div>

            <div class="error-message" id="errorMessage">
                <h3>❌ Processing Failed</h3>
                <p id="errorText">An error occurred during processing.</p>
            </div>

            <button class="download-btn" id="downloadBtn">
                💾 Download Word Document
            </button>
        </div>

        <script>
            let selectedSection = null;
            let processedFile = null;

            // Elements
            const englishSection = document.getElementById('englishSection');
            const mathSection = document.getElementById('mathSection');
            const dropZone = document.getElementById('dropZone');
            const fileInput = document.getElementById('fileInput');
            const progressSection = document.getElementById('progressSection');
            const progressFill = document.getElementById('progressFill');
            const statusText = document.getElementById('statusText');
            const successMessage = document.getElementById('successMessage');
            const errorMessage = document.getElementById('errorMessage');
            const downloadBtn = document.getElementById('downloadBtn');
            const apiKeyInput = document.getElementById('apiKeyInput');
            const apiKeyStatus = document.getElementById('apiKeyStatus');

            // Event listeners
            dropZone.addEventListener('click', () => {
                if (selectedSection) {
                    fileInput.click();
                } else {
                    alert('Please select a section first (English or Math)');
                }
            });
            fileInput.addEventListener('change', handleFileSelect);
            downloadBtn.addEventListener('click', downloadResult);

            // Drag and drop events
            ['dragenter', 'dragover', 'dragleave', 'drop'].forEach(eventName => {
                dropZone.addEventListener(eventName, preventDefaults, false);
            });

            dropZone.addEventListener('dragenter', () => dropZone.classList.add('dragover'));
            dropZone.addEventListener('dragleave', () => dropZone.classList.remove('dragover'));
            dropZone.addEventListener('drop', handleDrop);

            function preventDefaults(e) {
                e.preventDefault();
                e.stopPropagation();
            }

            function selectSection(section) {
                selectedSection = section;
                
                // Update UI
                englishSection.classList.remove('active');
                mathSection.classList.remove('active');
                
                if (section === 'english') {
                    englishSection.classList.add('active');
                    dropZone.querySelector('.drop-subtext').textContent = 'English section: Text processing, watermark removal & SAT detection';
                } else if (section === 'math') {
                    mathSection.classList.add('active');
                    dropZone.querySelector('.drop-subtext').textContent = 'Math section: LaTeX conversion, math detection & SAT processing';
                }
                
                // Enable drop zone
                dropZone.style.cursor = 'pointer';
                dropZone.style.opacity = '1';
            }

            function handleDrop(e) {
                dropZone.classList.remove('dragover');
                if (!selectedSection) {
                    alert('Please select a section first (English or Math)');
                    return;
                }
                const files = Array.from(e.dataTransfer.files);
                if (files.length > 0) {
                    handleFile(files[0]);
                }
            }

            function handleFileSelect(e) {
                if (!selectedSection) {
                    alert('Please select a section first (English or Math)');
                    return;
                }
                const files = Array.from(e.target.files);
                if (files.length > 0) {
                    handleFile(files[0]);
                }
            }

            function handleFile(file) {
                if (file.type !== 'application/pdf') {
                    showError('Only PDF files are supported');
                    return;
                }
                
                processFile(file);
            }

            async function processFile(file) {
                hideMessages();
                progressSection.style.display = 'block';
                dropZone.classList.add('processing');

                try {
                    updateProgress(10, 'Uploading file...');
                    
                    // Create FormData
                    const formData = new FormData();
                    formData.append('file', file);
                    formData.append('section_type', selectedSection);
                    
                    updateProgress(30, 'Processing document...');
                    
                    // Send to server
                    const response = await fetch('/convert', {
                        method: 'POST',
                        body: formData
                    });
                    
                    const result = await response.json();
                    
                    if (result.success) {
                        updateProgress(90, 'Generating Word document...');
                        processedFile = result.output_file;
                        showSuccess();
                    } else {
                        throw new Error(result.error || 'Processing failed');
                    }
                    
                } catch (error) {
                    console.error('Processing error:', error);
                    showError(error.message);
                } finally {
                    dropZone.classList.remove('processing');
                }
            }

            function updateProgress(percentage, text) {
                if (percentage !== null) {
                    progressFill.style.width = percentage + '%';
                }
                if (text) {
                    statusText.textContent = text;
                }
            }

            function showSuccess() {
                progressSection.style.display = 'none';
                successMessage.style.display = 'block';
                downloadBtn.style.display = 'inline-block';
                updateProgress(100, 'Complete!');
            }

            function showError(message) {
                progressSection.style.display = 'none';
                errorMessage.style.display = 'block';
                document.getElementById('errorText').textContent = message;
            }

            function hideMessages() {
                successMessage.style.display = 'none';
                errorMessage.style.display = 'none';
                downloadBtn.style.display = 'none';
            }

            async function updateApiKey() {
                const apiKey = apiKeyInput.value.trim();
                if (!apiKey) {
                    showApiKeyStatus('Please enter an API key', 'error');
                    return;
                }

                try {
                    const response = await fetch('/api-key', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({
                            api_key: apiKey,
                            provider: 'gemini'
                        })
                    });

                    const result = await response.json();
                    
                    if (result.success) {
                        showApiKeyStatus(result.message, 'success');
                        apiKeyInput.value = '';
                    } else {
                        showApiKeyStatus(result.error || 'Failed to update API key', 'error');
                    }
                } catch (error) {
                    console.error('API key update error:', error);
                    showApiKeyStatus('Failed to update API key. Please try again.', 'error');
                }
            }

            function showApiKeyStatus(message, type) {
                apiKeyStatus.textContent = message;
                apiKeyStatus.className = `api-key-status ${type}`;
                
                // Clear status after 5 seconds
                setTimeout(() => {
                    apiKeyStatus.textContent = '';
                    apiKeyStatus.className = 'api-key-status';
                }, 5000);
            }

            // AI Training Functions
            function startTraining() {
                const status = document.getElementById('trainingStatus');
                status.innerHTML = '🤖 Starting AI training session...<br>Upload a PDF to begin interactive training.';
                
                // Enable training mode
                window.trainingMode = true;
                showTrainingStatus('Training mode activated. Upload a PDF to provide feedback.', 'info');
            }

            function provideFeedback() {
                const feedback = prompt('Please provide feedback on the AI\'s performance:\n\n1. Was watermark detection accurate?\n2. Was the SAT format correctly identified?\n3. Any specific improvements needed?');
                
                if (feedback) {
                    // Send feedback to backend
                    fetch('/api/feedback', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({
                            feedback: feedback,
                            timestamp: new Date().toISOString()
                        })
                    })
                    .then(response => response.json())
                    .then(data => {
                        showTrainingStatus('✅ Feedback submitted successfully! Thank you for helping improve the AI.', 'success');
                    })
                    .catch(error => {
                        showTrainingStatus('❌ Failed to submit feedback. Please try again.', 'error');
                    });
                }
            }

            function showTrainingStatus(message, type) {
                const status = document.getElementById('trainingStatus');
                status.innerHTML = message;
                status.className = `training-status ${type}`;
                
                // Clear status after 5 seconds
                setTimeout(() => {
                    status.innerHTML = '';
                    status.className = 'training-status';
                }, 5000);
            }

            function downloadResult() {
                if (processedFile) {
                    window.open(`/download/${processedFile}`, '_blank');
                }
            }



            // Initialize with English section selected
            selectSection('english');
        </script>
    </body>
    </html>
    """
    
    return render_template_string(html_content)

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy', 
        'service': 'PDF Watermark Remover - Two Sections',
        'python_version': sys.version,
        'timestamp': time.time()
    })

@app.route('/convert', methods=['POST'])
def convert_document():
    """Convert PDF to Word document with watermark removal"""
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Get section type
        section_type = request.form.get('section_type', 'english')
        if section_type not in ['english', 'math', 'sat']:
            return jsonify({'error': 'Invalid section type'}), 400
        
        # Validate file
        if not file or not allowed_file(file.filename):
            return jsonify({'error': 'Only PDF files are supported'}), 400
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > MAX_FILE_SIZE:
            return jsonify({'error': f'File too large. Maximum size: {MAX_FILE_SIZE // (1024*1024)}MB'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        logger.info(f"File uploaded: {filename} ({file_size // 1024} KB) for {section_type} section")
        
        # Process document
        output_path, metadata = get_processor().process_document(file_path, section_type)
        
        # Clean up uploaded file
        try:
            os.remove(file_path)
            logger.info(f"Cleaned up uploaded file: {filename}")
        except:
            pass
        
        # Return success response
        return jsonify({
            'success': True,
            'message': f'Document converted successfully using {section_type} section',
            'output_file': os.path.basename(output_path),
            'metadata': metadata
        })
        
    except Exception as e:
        logger.error(f"Conversion failed: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500







@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download converted file"""
    try:
        file_path = os.path.join(PROCESSED_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Error handlers for production
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal server error: {error}")
    return jsonify({'error': 'Internal server error'}), 500

@app.errorhandler(Exception)
def handle_exception(e):
    logger.error(f"Unhandled exception: {str(e)}")
    return jsonify({'error': 'An unexpected error occurred'}), 500

@app.route('/train', methods=['POST'])
def train_model():
    """Train the watermark detection model with user feedback"""
    try:
        data = request.get_json()
        
        # Get user feedback
        user_feedback = {
            'remove': data.get('remove', []),
            'preserve': data.get('preserve', []),
            'document_path': data.get('document_path', ''),
            'comment': data.get('comment', ''),
            'action': data.get('action', ''),  # 'approve' or 'deny'
            'timestamp': datetime.now().isoformat()
        }
        
        # Train the model based on user action
        processor = DocumentProcessor()
        success = False
        
        if user_feedback['action'] == 'deny':
            # User denied the output, retrain and regenerate
            success = processor.watermark_remover.train_on_document(
                user_feedback['document_path'], 
                user_feedback
            )
            message = 'Model retrained based on your feedback. Please try processing again.'
        elif user_feedback['action'] == 'approve':
            # User approved, learn from the successful output
            success = processor.watermark_remover.learn_from_approval(
                user_feedback['document_path'], 
                user_feedback
            )
            message = 'Model learned from your approval. Thank you for the feedback!'
        else:
            message = 'Invalid action. Please use "approve" or "deny".'
        
        # Save feedback for analysis
        feedback_file = 'training_feedback.json'
        if os.path.exists(feedback_file):
            with open(feedback_file, 'r') as f:
                feedback_data = json.load(f)
        else:
            feedback_data = []
        
        feedback_data.append(user_feedback)
        
        with open(feedback_file, 'w') as f:
            json.dump(feedback_data, f, indent=2)
        
        return jsonify({
            'success': success,
            'message': message
        })
        
    except Exception as e:
        logger.error(f"Training error: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

if __name__ == '__main__':
    print("Starting PDF Watermark Remover - Enhanced SAT Processing...")
    print(f"Python {sys.version} detected")
    print("All dependencies loaded successfully")
    print("Watermark removal algorithms ready")
    print("SAT document processor ready")
    print("AI-powered enhancement ready (API key configurable via web interface)")
    
    # Set startup time for health checks
    app_startup_time = time.time()
    print(f"Startup timestamp: {app_startup_time}")
    
    # Get port from environment variable (for Railway/Heroku)
    port = int(os.environ.get('PORT', 5000))
    
    # Production optimizations for Railway
    is_production = os.environ.get('RAILWAY_ENVIRONMENT') or os.environ.get('RAILWAY_STATIC_URL')
    
    if is_production:
        print("🚀 Production mode detected - Railway deployment")
        host = '0.0.0.0'
        debug = False
        threaded = True
    else:
        print("🔧 Development mode detected - Local testing")
        host = 'localhost'
        debug = True
        threaded = False
    
    print(f"\nServer will start at: http://{host}:{port}")
    print("Two sections: English (text + SAT) and Math (LaTeX + SAT)")
    print("SAT documents auto-detected in both sections")
    print("Health check at: http://{host}:{port}/health")
    print(f"Production mode: {is_production}")
    print(f"Debug mode: {debug}")
    print(f"Threaded: {threaded}")
    print("Press Ctrl+C to stop the server")
    print("\n" + "="*60)
    
    try:
        app.run(
            debug=debug, 
            host=host, 
            port=port,
            threaded=threaded
        )
    except KeyboardInterrupt:
        print("\nServer stopped by user")
    except Exception as e:
        print(f"\nServer error: {e}")
        sys.exit(1)
