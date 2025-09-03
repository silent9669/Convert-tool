#!/usr/bin/env python3
"""
Final Optimized Digital SAT PDF Processor
Addresses formatting and choice detection issues
"""

import os
import re
import fitz
import logging
import json
import time
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from PIL import Image
import io
import base64
from collections import Counter, defaultdict

# Try to import advanced dependencies
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Import local configuration
try:
    from local_config import GEMINI_API_KEY, TEST_MODE, DEBUG_LOGGING
    LOCAL_CONFIG_AVAILABLE = True
except ImportError:
    LOCAL_CONFIG_AVAILABLE = False
    GEMINI_API_KEY = None
    TEST_MODE = False
    DEBUG_LOGGING = False

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FinalOptimizedProcessor:
    """Final optimized Digital SAT processor with corrected formatting"""
    
    def __init__(self, gemini_api_key: str = None):
        # Use local config first, then parameter, then environment variable
        if LOCAL_CONFIG_AVAILABLE and GEMINI_API_KEY:
            self.gemini_api_key = GEMINI_API_KEY
        else:
            self.gemini_api_key = gemini_api_key or os.environ.get('GEMINI_API_KEY')
        
        self.setup_ai()
        self.setup_optimized_patterns()
        self.setup_logging()
        
    def setup_ai(self):
        """Setup AI models for question detection"""
        if GEMINI_AVAILABLE and self.gemini_api_key:
            try:
                genai.configure(api_key=self.gemini_api_key)
                self.model = genai.GenerativeModel('gemini-2.0-flash-exp')
                self.ai_available = True
                logger.info("✅ Gemini AI configured successfully")
            except Exception as e:
                logger.warning(f"⚠️ Gemini AI setup failed: {e}")
                self.ai_available = False
        else:
            self.ai_available = False
    
    def setup_logging(self):
        """Setup logging for debugging"""
        logging.basicConfig(level=logging.INFO)
    
    def setup_optimized_patterns(self):
        """Setup optimized patterns based on analysis"""
        
        self.patterns = {
            # Module headers
            'module_header': [
                r'^MODULE\s+\d+$',
                r'^Module\s+\d+$',
                r'^Section\s+\d+$',
            ],
            
            # Question number patterns - FIXED for proper detection
            'question_number': [
                r'^Q\d+\.\s*$',  # "Q1." - Standard format
                r'^\d+\.\s*$',  # "1." with optional spaces
                r'^Question\s+\d+\.?\s*$',  # "Question 1" or "Question 1."
                r'\d+\s+\[Q\s+Mark\s+for\s+Review',  # "2 [Q Mark for Review" - OCR pattern
                r'\d+\s+Mark\s+for\s+Review',  # "1 Mark for Review" - OCR pattern
                r'^\d+\s*$',  # Just a number on its own line
                r'^\d+\s+\[',  # "2 [" - OCR pattern
                r'^\d+\s+\[Q',  # "2 [Q" - OCR pattern
            ],
            
            # Multiple choice patterns - FIXED for proper detection
            'multiple_choice': [
                r'^[A-D]\.\s+',  # "A. Text"
                r'^[A-D]\)\s+',  # "A) Text"
                r'^[A-D]\s+[A-Za-z]',  # "A Text" (no punctuation)
            ],
            
            # Passage detection
            'passage_start': [
                r'^[A-Z][a-z]+(?:\s+[a-z]+){4,}',  # 5+ word sentences starting with capital
                r'^The\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # "The unique subak water..."
                r'^Until\s+\d{4}',  # "Until 1917..."
                r'^[A-Z][a-z]+\s+[a-z]+-[a-z]+',  # "The Egyptian plover-a bird..."
                r'^[A-Z][a-z]+\s+[a-z]+\s+\([^)]+\)',  # "The mihrab (or niche)..."
                r'^[A-Z][a-z]+\s+[a-z]+\s+[a-z]+\s+[a-z]+',  # 4+ word sentences
                r'^Some\s+[a-z]+\s+[a-z]+',  # "Some ethicists challenge..."
                r'^As\s+[A-Z][a-z]+\s+[a-z]+',  # "As Rachana Kamtekar observes..."
                r'^In\s+[A-Z][a-z]+\s+[a-z]+',  # "In recent years..."
                r'^For\s+[a-z]+\s+[a-z]+',  # "For many years..."
            ],
            
            # Question content patterns
            'question_content': [
                r'Which\s+choice\s+completes\s+the\s+text',
                r'What\s+is\s+the\s+value\s+of',
                r'Based\s+on\s+the\s+text',
                r'According\s+to\s+the\s+passage',
                r'The\s+author\s+implies',
                r'In\s+the\s+context',
                r'Which\s+choice\s+most\s+effectively',
                r'Which\s+choice\s+best\s+completes',
                r'What\s+does\s+the\s+author',
                r'According\s+to\s+the\s+graph',
                r'Based\s+on\s+the\s+table',
                r'Which\s+of\s+the\s+following',
            ],
            
            # Enhanced watermark patterns
            'watermark': [
                r'^A-\s*A\+\s*\d{2}:\d{2}:\d{2}$',  # "A- A+ 21:23:57"
                r'^\s*[-+]+\s*$',  # Lines with just dashes or pluses
                r'^\s*[A-Z]+\s*$',  # Lines with just capital letters
                r'^\s*\d+\s*$',  # Lines with just numbers (if standalone)
                r'^\s*[A-Z][a-z]+\s+\d{2}:\d{2}:\d{2}\s*$',  # "Time 21:23:57"
                r'^\s*[A-Z][a-z]+\s+\d{2}:\d{2}:\d{2}\s*[A-Z][a-z]+\s*$',  # "Time 21:23:57 Date"
                r'^\s*[A-Z][a-z]+\s+\d{2}:\d{2}:\d{2}\s*[A-Z][a-z]+\s+\d{2}:\d{2}:\d{2}\s*$',  # "Time 21:23:57 Date 21:23:57"
                r'^\(Ver\s+[A-Z]\)$',  # "(Ver C)"
                r'^https://t\.me/',  # Telegram links
                r'^@\w+$',  # Social media handles
                r'^\d+\.\d+\.\d+\.\d+',  # IP addresses
                r'^\d+\s*:\d+\s*in\s+a$',  # "31:44 in a"
                r'^Directions\s+v\s+Hide',  # "Directions v Hide"
                r'^More$',  # Just "More"
                r'^bps:\d+kb',  # "bps:2kb"
                r'^fps:\d+',  # "fps:26"
                r'^frame\s+size:\d+kb',  # "frame size:Okb"
            ],
        }
    
    def process_digital_sat_pdf(self, pdf_path: str) -> Dict:
        """Main method to process Digital SAT PDF with optimized algorithms"""
        try:
            logger.info(f"Processing Digital SAT PDF: {pdf_path}")
            
            # Extract text and images using enhanced OCR
            text_content, images = self._extract_content_with_ocr(pdf_path)
            
            # Clean and structure the content
            cleaned_text = self._optimized_text_cleaning(text_content)
            structure = self._optimized_structure_detection(cleaned_text)
            
            # Enhance with AI if available
            if self.ai_available:
                try:
                    structure = self._enhance_with_ai_optimized(structure, cleaned_text)
                except Exception as e:
                    logger.warning(f"AI enhancement skipped: {e}")
            
            # Process images
            processed_images = self._process_images(images, structure)
            
            return {
                'text': cleaned_text,
                'structure': structure,
                'images': processed_images,
                'metadata': {
                    'total_questions': len(structure.get('questions', [])),
                    'total_passages': len(structure.get('passages', [])),
                    'sections': structure.get('sections', []),
                    'processing_time': time.time()
                }
            }
                
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise
    
    def _extract_content_with_ocr(self, pdf_path: str) -> Tuple[str, List]:
        """Extract content with OCR"""
        try:
            doc = fitz.open(pdf_path)
            text_parts = []
            images = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # First try regular text extraction
                text = page.get_text()
                
                # Enhanced OCR detection and processing
                if len(text.strip()) < 200:  # Lower threshold for better detection
                    logger.info(f"Page {page_num + 1} has little text, attempting OCR...")
                    if OCR_AVAILABLE:
                        # Convert page to image for OCR
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Higher resolution
                        img_data = pix.tobytes("png")
                        img_pil = Image.open(io.BytesIO(img_data))
                        
                        # Enhanced OCR with multiple configurations
                        ocr_configs = [
                            '--psm 6',  # Uniform block of text
                            '--psm 3',  # Fully automatic page segmentation
                            '--psm 4',  # Assume a single column of text
                        ]
                        
                        best_text = ""
                        for config in ocr_configs:
                            try:
                                ocr_text = pytesseract.image_to_string(img_pil, config=config)
                                if len(ocr_text.strip()) > len(best_text.strip()):
                                    best_text = ocr_text
                            except Exception as e:
                                logger.warning(f"OCR config {config} failed: {e}")
                        
                        if best_text.strip():
                            text = best_text
                            logger.info(f"OCR extracted {len(text)} characters from page {page_num + 1}")
                    
                if text.strip():
                    text_parts.append(text)
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            img_pil = Image.open(io.BytesIO(img_data))
                            
                            images.append({
                                'page': page_num + 1,
                                'index': img_index,
                                'image': img_pil,
                                'bbox': img[1:5],  # x0, y0, x1, y1
                                'description': f"Page {page_num + 1} Image {img_index + 1}",
                                'position': {
                                    'x': img[1],
                                    'y': img[2],
                                    'width': img[3] - img[1],
                                    'height': img[4] - img[2]
                                }
                            })
                        
                        pix = None  # Free memory
                    except Exception as e:
                        logger.warning(f"Image extraction failed: {e}")
            
            doc.close()
            return '\n'.join(text_parts), images
            
        except Exception as e:
            logger.error(f"Content extraction failed: {e}")
            raise
    
    def _optimized_text_cleaning(self, text: str) -> str:
        """Optimized text cleaning"""
        if not text:
            return ""
            
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Enhanced watermark detection
            if self._is_watermark(line):
                continue
            
            # Clean up common artifacts
            line = re.sub(r'\s+', ' ', line)  # Multiple spaces to single
            line = re.sub(r'[^\w\s\.\,\;\:\!\?\(\)\[\]\-\+\=\<\>]', '', line)  # Remove special chars
            
            if line.strip() and len(line.strip()) > 2:  # Only keep substantial lines
                cleaned_lines.append(line.strip())
        
        return '\n'.join(cleaned_lines)
    
    def _is_watermark(self, line: str) -> bool:
        """Watermark detection"""
        line_clean = line.strip()
        
        # Check exact patterns first
        for pattern in self.patterns['watermark']:
            if re.search(pattern, line_clean):
                return True
        
        # Enhanced heuristic checks
        if len(line_clean) < 3:  # Very short lines
            return True
        
        if re.match(r'^[A-Z\s\d\-\+]+$', line_clean):  # Only caps, numbers, dashes, pluses
            return True
        
        if re.match(r'^\d+\s*$', line_clean):  # Just a number
            return True
        
        return False
    
    def _optimized_structure_detection(self, text: str) -> Dict:
        """Optimized structure detection with proper formatting"""
        lines = text.split('\n')
        
        structure = {
            'passages': [],
            'questions': [],
            'math_functions': [],
            'sections': [],
            'module_headers': [],
            'total_questions': 0
        }
        
        current_passage = None
        current_question = None
        question_counter = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check for module headers
            if self._matches_pattern(line, self.patterns['module_header']):
                structure['module_headers'].append({
                    'text': line,
                    'line_number': i,
                    'type': 'module'
                })
                continue
            
            # FIXED: Enhanced question detection with proper numbering
            if self._is_question_number(line):
                # Save previous question if exists
                if current_question:
                    structure['questions'].append(current_question)
                    question_counter += 1
                
                # Start new question with proper numbering
                question_num = self._extract_question_number(line)
                current_question = {
                    'number': question_num,
                    'text': f"Question {question_num}",  # FIXED: Proper format
                    'content': [],
                    'choices': [],
                    'line_number': i,
                    'type': 'unknown',
                    'passage_content': [],
                    'images': []
                }
                continue
                
            # FIXED: Enhanced multiple choice detection
            if self._is_multiple_choice(line):
                if current_question:
                    choice = self._parse_choice(line)
                    if choice:
                        current_question['choices'].append(choice)
                continue
            
            # Enhanced passage detection
            if self._is_passage_start(line):
                if current_passage:
                    structure['passages'].append(current_passage)
                
                current_passage = {
                    'text': line,
                    'content': [line],
                    'line_number': i,
                    'type': 'reading_passage'
                }
                continue
                
            # Enhanced content association
            if current_question:
                # Add content to questions with better logic
                if self._is_question_content(line):
                    current_question['content'].append(line)
                elif self._is_passage_content(line):
                    current_question['passage_content'].append(line)
            elif current_passage:
                # Only add substantial content to passages
                if len(line.split()) >= 3 and len(line) > 20:
                    current_passage['content'].append(line)
        
        # Add final passage and question
        if current_passage:
            structure['passages'].append(current_passage)
        if current_question:
            structure['questions'].append(current_question)
            question_counter += 1
        
        structure['total_questions'] = question_counter
        
        return structure
    
    def _is_question_number(self, line: str) -> bool:
        """Question number detection - FIXED"""
        # Main pattern for SAT format
        if re.match(r'^Q\d+\.\s*$', line):  # "Q1."
            return True
        
        # Other patterns including OCR-specific patterns
        patterns = [
            r'^\d+\.\s*$',  # "1." with optional spaces
            r'^Question\s+\d+\.?\s*$',  # "Question 1" or "Question 1."
            r'\d+\s+\[Q\s+Mark\s+for\s+Review',  # "2 [Q Mark for Review" - OCR pattern
            r'\d+\s+Mark\s+for\s+Review',  # "1 Mark for Review" - OCR pattern
            r'^\d+\s*$',  # Just a number on its own line
            r'^\d+\s+\[',  # "2 [" - OCR pattern
            r'^\d+\s+\[Q',  # "2 [Q" - OCR pattern
        ]
        
        for pattern in patterns:
            if re.search(pattern, line):  # Use search instead of match for OCR patterns
                return True
        
        return False
    
    def _is_multiple_choice(self, line: str) -> bool:
        """Multiple choice detection - FIXED"""
        patterns = [
            r'^[A-D]\.\s+',  # "A. Text"
            r'^[A-D]\)\s+',  # "A) Text"
            r'^[A-D]\s+[A-Za-z]',  # "A Text" (no punctuation)
        ]
        
        for pattern in patterns:
            if re.search(pattern, line):
                return True
        
        return False
    
    def _is_passage_start(self, line: str) -> bool:
        """Passage start detection"""
        # Must be substantial text
        if len(line.split()) < 5:
            return False
        
        # Must start with capital letter
        if not line[0].isupper():
            return False
        
        # Must not end with period (likely a question)
        if line.endswith('.'):
            return False
        
        # Must not be a question pattern
        if '?' in line:
            return False
        
        # Check against enhanced passage patterns
        return self._matches_pattern(line, self.patterns['passage_start'])
    
    def _is_question_content(self, line: str) -> bool:
        """Check if line is question content"""
        return self._matches_pattern(line, self.patterns['question_content'])
    
    def _is_passage_content(self, line: str) -> bool:
        """Check if line is passage content"""
        # Must be substantial text
        if len(line.split()) < 10:
            return False
        
        # Must start with capital letter
        if not line[0].isupper():
            return False
        
        # Must not be a question
        if '?' in line:
            return False
        
        return True
    
    def _matches_pattern(self, text: str, patterns: List[str]) -> bool:
        """Check if text matches any pattern in the list"""
        for pattern in patterns:
            if re.search(pattern, text):
                return True
        return False
    
    def _extract_question_number(self, text: str) -> str:
        """Extract question number from text - FIXED"""
        # Handle Q1. format
        match = re.search(r'Q(\d+)', text)
        if match:
            return match.group(1)
        
        # Handle OCR patterns like "2 [Q Mark for Review" or "1 Mark for Review"
        match = re.search(r'(\d+)\s+\[Q\s+Mark\s+for\s+Review', text)
        if match:
            return match.group(1)
        
        match = re.search(r'(\d+)\s+Mark\s+for\s+Review', text)
        if match:
            return match.group(1)
        
        # Handle regular number format
        match = re.search(r'(\d+)', text)
        return match.group(1) if match else text
    
    def _parse_choice(self, text: str) -> Optional[Dict]:
        """Enhanced multiple choice option parsing - FIXED"""
        if not text:
            return None
        
        # Extract option letter - FIXED patterns
        option_match = re.match(r'^([A-D])[\)\.\s]*', text)
        if not option_match:
            return None
        
        option = option_match.group(1)
        choice_text = re.sub(r'^[A-D][\)\.\s]*', '', text).strip()
        
        if choice_text and len(choice_text) > 2:  # Must have substantial text
            return {
                'option': option,
                'text': choice_text,
                'is_correct': False
            }
        
        return None
    
    def _enhance_with_ai_optimized(self, structure: Dict, text: str) -> Dict:
        """Optimized AI enhancement"""
        try:
            if not self.ai_available:
                return structure

            # Create optimized prompt for AI
            prompt = f"""
            Analyze this Digital SAT content and improve the structure detection:
            
            Text sample: {text[:2000]}...
            
            Current structure: {json.dumps(structure, indent=2)[:800]}...
            
            Based on the expected SAT format:
            - Questions should have "Question X" headers (where X is the number)
            - Each question should have A, B, C, D multiple choice options
            - Passage content should be separated from question content
            - Remove any watermarks or technical artifacts
            
            Please provide improvements in JSON format for:
            1. Better question detection and numbering
            2. Proper passage separation
            3. Multiple choice option parsing
            4. Content organization
            
            Focus on accuracy and proper SAT format structure.
            Return only valid JSON.
            """
            
            response = self.model.generate_content(prompt)
            
            if response and response.text:
                try:
                    # Try to parse AI response
                    ai_enhancement = json.loads(response.text)
                    
                    # Merge AI improvements with existing structure
                    structure = self._merge_ai_enhancements(structure, ai_enhancement)
                    
                except json.JSONDecodeError:
                    logger.warning("AI response not valid JSON")

        except Exception as e:
            logger.warning(f"AI enhancement failed: {e}")
        
        return structure
    
    def _merge_ai_enhancements(self, original: Dict, ai_enhancement: Dict) -> Dict:
        """Merge AI enhancements with original structure"""
        # Deep merge AI improvements
        for key, value in ai_enhancement.items():
            if key in original and isinstance(original[key], list) and isinstance(value, list):
                # Merge lists
                original[key].extend(value)
            elif key in original and isinstance(original[key], dict) and isinstance(value, dict):
                # Recursively merge dicts
                original[key] = self._merge_ai_enhancements(original[key], value)
            else:
                # Replace or add new values
                original[key] = value
        
        return original
    
    def _process_images(self, images: List, structure: Dict) -> List:
        """Process and organize images"""
        processed_images = []
        
        for img in images:
            # Try to associate images with questions or passages
            associated_content = self._find_image_association(img, structure)
            
            processed_img = {
                'page': img['page'],
                'index': img['index'],
                'image': img['image'],
                'bbox': img['bbox'],
                'description': img['description'],
                'position': img.get('position', {}),
                'associated_content': associated_content
            }
            
            processed_images.append(processed_img)
            
            # Add to structure if associated
            if associated_content:
                if associated_content['type'] == 'question':
                    for question in structure['questions']:
                        if question['number'] == associated_content['id']:
                            question['images'].append(processed_img)
                            break
        
        return processed_images
    
    def _find_image_association(self, image: Dict, structure: Dict) -> Optional[Dict]:
        """Image association logic"""
        # Use position information for better association
        img_page = image['page']
        img_position = image.get('position', {})
        
        # Find nearby questions based on page and position
        for question in structure['questions']:
            question_line = question.get('line_number', 0)
            
            # Rough estimation: associate images with questions on same page
            if question_line // 50 == img_page - 1:  # Rough page estimation
                return {
                    'type': 'question',
                    'id': question['number']
                }
        
        return None
    
    def generate_word_document(self, structure: Dict, images: List, output_path: str) -> str:
        """Generate Word document from processed structure - FIXED formatting"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Add title
            title = doc.add_heading('Digital SAT Practice Test', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add module header
            if structure.get('module_headers'):
                module_header = doc.add_heading('MODULE 1', level=1)
                module_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process questions with FIXED formatting
            for question in structure.get('questions', []):
                # Add question number with proper format
                question_para = doc.add_paragraph()
                question_para.add_run(f"Question {question['number']}").bold = True
                doc.add_paragraph()  # Spacing
                
                # Add images if present
                if 'images' in question and question['images']:
                    for img in question['images']:
                        try:
                            # Convert PIL image to bytes
                            img_bytes = io.BytesIO()
                            img['image'].save(img_bytes, format='PNG')
                            img_bytes.seek(0)
                            
                            # Add image to document
                            doc.add_picture(img_bytes, width=Inches(4))
                            doc.add_paragraph()  # Spacing
                        except Exception as e:
                            logger.warning(f"Image insertion failed: {e}")
                
                # Add passage content if present
                if 'passage_content' in question and question['passage_content']:
                    for passage_line in question['passage_content']:
                        if passage_line.strip():
                            doc.add_paragraph(passage_line.strip())
                    doc.add_paragraph()  # Spacing
                
                # Add question content
                for content_line in question.get('content', []):
                    if content_line.strip():
                        doc.add_paragraph(content_line.strip())
                
                # Add multiple choice options with FIXED formatting
                if question.get('choices'):
                    for choice in question['choices']:
                        choice_text = f"{choice['option']}. {choice['text']}"
                        doc.add_paragraph(choice_text)
                
                doc.add_paragraph()  # Spacing between questions
            
            # Save document
            doc.save(output_path)
            logger.info(f"Word document generated: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Word document generation failed: {e}")
            raise

# Test the final optimized processor
if __name__ == "__main__":
    processor = FinalOptimizedProcessor()
    
    # Test with all 3 PDF files
    test_files = [
        'rss/input 1.pdf',
        'rss/input 2.pdf', 
        'rss/easy_testing.pdf'
    ]
    
    for pdf_file in test_files:
        try:
            print(f"\n=== TESTING {pdf_file} ===")
            result = processor.process_digital_sat_pdf(pdf_file)
            
            print(f"Questions detected: {result['metadata']['total_questions']}")
            print(f"Passages detected: {result['metadata']['total_passages']}")
            
            if result['structure']['questions']:
                print("First few questions:")
                for i, q in enumerate(result['structure']['questions'][:3]):
                    print(f"  Question {q['number']}: {q['text']}")
                    if q['choices']:
                        print(f"    Choices: {len(q['choices'])} options")
            
            # Generate Word document
            output_name = f"output/final_{Path(pdf_file).stem}_optimized.docx"
            processor.generate_word_document(result['structure'], result['images'], output_name)
            print(f"Generated: {output_name}")
            
        except Exception as e:
            print(f"Error processing {pdf_file}: {e}")
